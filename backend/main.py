# main.py
# Agentic Framework Backend for the AI Business Analyst - FINAL

from flask import Flask, request, jsonify, redirect, send_file, send_from_directory
from flask_cors import CORS
import os
# Ensure Flask does not auto-load .env with default UTF-8 decoding
os.environ.setdefault("FLASK_SKIP_DOTENV", "1")
from io import BytesIO
import requests
import json
import uuid
import base64
import io
import re
from datetime import datetime, timezone
from werkzeug.utils import secure_filename
from sqlalchemy import text

# --- File Parsing Libraries ---
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
import docx

# --- Azure SDKs ---
from azure.communication.email import EmailClient

# --- Database imports ---
from database import (
    init_db, get_db, save_document_to_db, save_document_to_db_direct, save_analysis_to_db,
    add_to_vector_db, search_vector_db, Document, Analysis,
    save_approval_to_db, get_approval_from_db, update_approval_in_db, update_approval_in_db_with_data,
    get_all_documents_from_db, get_all_documents_from_db_direct, get_all_analyses_from_db, get_analysis_details_from_db,
    check_document_exists_by_name, check_document_exists_by_name_direct, get_document_by_id, get_analysis_by_id_from_db,
    delete_from_vector_db
)

# --- Multi-User Database imports ---
from database_multi_user import (
    User, Project, TextInput, ProjectStatus, ProjectUserRole,
    create_project, get_user_projects, add_user_to_project, save_text_input,
    project_documents
)

# --- Authentication imports ---
from auth_system import (
    auth_manager, token_required, role_required, project_access_required,
    register_user, login_user, get_user_profile, update_user_profile,
    change_password, logout_user, get_all_users, update_user_status
)

# --- Integration imports ---
from integration_services import IntegrationManager

# --- Enhanced Document Generation ---
from enhanced_document_generator import EnhancedDocumentGenerator

# --- Configuration imports ---
from config import (
    GEMINI_API_URL, ACS_CONNECTION_STRING, ACS_SENDER_ADDRESS,
    APPROVAL_RECIPIENT_EMAIL, BACKEND_BASE_URL, ADO_ORGANIZATION_URL,
    ADO_PROJECT_NAME, ADO_PAT, DATABASE_URL, GEMINI_API_KEY,
    LUCID_ENABLED
)

# Initialize the Flask application
app = Flask(__name__)

# Initialize database
init_db()

# --- Configuration ---
CORS(app, resources={r"/api/*": {"origins": "*"}})

# Initialize integration manager
integration_manager = IntegrationManager()

# Initialize enhanced document generator
enhanced_doc_generator = EnhancedDocumentGenerator(None)  # Will be set with LLM engine

# --- In-memory storage for approvals (fallback only) ---
approval_statuses = {}

# --- Token Consumption Tracking ---
token_consumption_logs = []

def log_token_consumption(stage, tokens_used, model_used="gemini-pro", details=None):
    """Log token consumption for each stage with detailed information"""
    log_entry = {
        "timestamp": datetime.now().isoformat(),
        "stage": stage,
        "tokens_used": tokens_used,
        "model_used": model_used,
        "details": details or {}
    }
    token_consumption_logs.append(log_entry)
    
    # Save to file for persistence
    try:
        with open("token_logs.json", "w") as f:
            json.dump(token_consumption_logs, f, indent=2)
    except Exception as e:
        print(f"Warning: Could not save token logs: {e}")
    
    # Print detailed log to console
    print(f"🔍 TOKEN LOG: {stage.upper()}")
    print(f"   📊 Tokens Used: {tokens_used:,}")
    print(f"   🤖 Model: {model_used}")
    print(f"   ⏰ Timestamp: {log_entry['timestamp']}")
    if details:
        print(f"   📝 Details: {details}")
    print(f"   📈 Total Tokens So Far: {sum(log['tokens_used'] for log in token_consumption_logs):,}")
    print("─" * 50)

# ----------------------------------------------------------------------------
# Line of Business (LOB) classification
# ----------------------------------------------------------------------------
def classify_line_of_business(text: str) -> dict:
    """Classify the uploaded document's line of business and region.

    Focus: P&C Insurance for US and Europe. Uses simple keyword heuristics.
    Returns a dict with: industry, segment, region, confidence, matched_terms.
    """
    if not text:
        return {
            "industry": "unknown",
            "segment": "unknown",
            "region": "unknown",
            "confidence": 0.0,
            "matched_terms": []
        }

    lowered = text.lower()

    insurance_terms = [
        "insurance", "insurer", "policy", "premium", "claim", "underwriting",
        "broker", "reinsurance", "adjuster", "loss ratio", "actuarial", "coverage",
        "deductible", "endorsement", "exposure", "policyholder"
    ]
    pnc_terms = [
        "property", "casualty", "auto", "motor", "homeowners", "commercial lines",
        "workers' compensation", "general liability", "umbrella", "marine", "aviation",
        "bond", "surety", "acord", "iso form"
    ]
    us_markers = [
        "naic", "acord", "state farm", "progressive", "geico", "usaa", "hipaa", "pci-dss"
    ]
    eu_markers = [
        "gdpr", "eiopa", "solvency ii", "lloyd's", "fca", "prudential regulation"
    ]

    # Sub-line of business (LoB) categories within P&C
    pnc_categories: dict[str, list[str]] = {
        "personal_auto": ["personal auto", "auto policy", "motor policy", "bodily injury", "collision", "comprehensive", "vehicle", "driver"],
        "commercial_auto": ["commercial auto", "fleet", "trucking", "motor carrier", "garage liability"],
        "homeowners": ["homeowners", "dwelling", "ho-3", "ho3", "renters", "condo", "hazard"],
        "property": ["property", "fire", "inland marine", "builders risk", "business interruption", "contents"],
        "general_liability": ["general liability", "cgl", "liability", "premises operations", "products completed operations"],
        "workers_comp": ["workers' compensation", "work comp", "wc", "experience mod", "ncci"],
        "environmental": ["environmental", "pollution", "epa", "spill", "remediation"],
        "professional_liability": ["professional liability", "e&o", "errors and omissions", "malpractice"],
        "financial_lines": ["d&o", "directors and officers", "crime", "fidelity", "kidnap", "ransom", "fi", "financial institutions"],
        "marine": ["ocean marine", "cargo", "hull", "inland marine"],
        "aviation": ["aviation", "aircraft", "hangar"],
        "cyber": ["cyber", "privacy", "data breach", "ransomware"],
        "umbrella": ["umbrella", "excess liability", "excess"]
    }

    matched = []
    score = 0
    for term in insurance_terms:
        if term in lowered:
            matched.append(term)
            score += 1
    for term in pnc_terms:
        if term in lowered:
            matched.append(term)
            score += 2

    region = "unknown"
    if any(m in lowered for m in us_markers):
        region = "US"
        score += 1
    if any(m in lowered for m in eu_markers):
        region = "Europe" if region == "unknown" else region
        score += 1

    industry = "Insurance" if score > 0 else "unknown"
    segment = "P&C" if any(t in lowered for t in pnc_terms) else ("Insurance" if industry == "Insurance" else "unknown")

    # Normalize region for insurance context
    if industry == "Insurance" and segment == "P&C" and region == "unknown":
        # Default to US/EU unknown if insurance P&C but no markers
        region = "US/EU-unknown"

    # Detect sub-LoB categories and assign scores
    category_hits = []
    for category, keywords in pnc_categories.items():
        cat_terms = []
        cat_score = 0
        for kw in keywords:
            if kw in lowered:
                cat_terms.append(kw)
                cat_score += 1
        if cat_terms:
            category_hits.append({
                "name": category,
                "score": cat_score,
                "matched_terms": cat_terms
            })

    # Normalize confidences for categories
    max_cat = max((c["score"] for c in category_hits), default=0)
    for c in category_hits:
        c["confidence"] = round(min(1.0, (c["score"] / max(1, max_cat)) * 0.9 + 0.1), 2)

    confidence = min(1.0, score / 12.0)

    return {
        "industry": industry,
        "segment": segment,
        "region": region,
        "confidence": round(confidence, 2),
        "matched_terms": matched[:20],
        "lob_categories": sorted(category_hits, key=lambda x: x["score"], reverse=True)[:6]
    }

@app.route("/api/token_logs", methods=['GET'])
def get_token_logs():
    """Get token consumption logs"""
    return jsonify(token_consumption_logs)

# --- Document Management ---
# --- Document Management ---
@app.route("/api/upload_document", methods=['POST'])
def upload_document():
    """Upload and store a document"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part in the request"}), 400

        file = request.files['file']
        print(f"📥 [Upload] Started for file: {file.filename}")

        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        if not file.filename.lower().endswith(('.pdf', '.docx')):
            return jsonify({"error": "Only PDF and DOCX files are allowed"}), 400

        print(f"DEBUG: Starting document upload for {file.filename}")

        # Check if document with same name already exists
        print(f"🔍 [Upload] Checking if file '{file.filename}' already exists...")
        if check_document_exists_by_name_direct(file.filename):
            print(f"⚠️ [Upload] File already exists: {file.filename}")
            print(f"DEBUG: Document with name '{file.filename}' already exists, skipping upload")
            return jsonify({
                "error": f"Document with name '{file.filename}' already exists. Please use a different name.",
                "duplicate": True
            }), 409  # Conflict status code

        # Read file content once and store it
        file_content_bytes = file.read()
        file_size = len(file_content_bytes)
        print(f"DEBUG: File size: {file_size} bytes")
        
        # Create uploads directory if it doesn't exist
        uploads_dir = "uploads"
        if not os.path.exists(uploads_dir):
            os.makedirs(uploads_dir)
            print(f"📁 Created uploads directory: {uploads_dir}")
        
        # Extract content from file
        file_content, images, error = agent_extract_content(io.BytesIO(file_content_bytes), file.filename)
        if error:
            print(f"DEBUG: Content extraction failed: {error}")
            return jsonify({"error": f"Failed to extract content: {error}"}), 500

        print(f"DEBUG: Content extracted successfully, length: {len(file_content)} characters")

        # Create document record
        doc_id = str(uuid.uuid4())
        # Now persist original upload to disk under generated ID
        file_path = f"{uploads_dir}/{doc_id}_{file.filename}"
        with open(file_path, 'wb') as f:
            f.write(file_content_bytes)
        print(f"💾 [File] Saved file to disk: {file_path}")
        
        # Classify LOB for the document
        lob_info = classify_line_of_business(file_content)

        document_data = {
            "id": doc_id,
            "name": file.filename,
            "uploadDate": datetime.now().isoformat(),
            "fileType": file.filename.split('.')[-1].lower(),
            "size": file_size,
            "status": "uploaded",
            "lob": lob_info
        }
        
        print(f"DEBUG: Document data created with ID: {doc_id}")
        
        # Save to database using SQLAlchemy, add to vector DB with LOB metadata
        db = next(get_db())
        try:
            print(f"💾 [DB] Saving document to database using SQLAlchemy: {file.filename}")
            # The line below is the one you need to check.
            saved_doc = save_document_to_db(db, user_email="guest", file_name=file.filename, file_type=document_data["fileType"], file_path=file_path, file_content=file_content, meta=document_data, status="uploaded")
            
            # --- START OF THE REQUIRED CHANGE ---
            if saved_doc is None:
                # If the database save failed, return an error.
                return jsonify({
                    "status": "error",
                    "message": "Failed to save document to database."
                }), 500
            # --- END OF THE REQUIRED CHANGE ---

            print(f"✅ [DB] Document stored with ID: {doc_id}")

            # Add to vector database
            print(f"📡 [VectorDB] Adding to vector DB...")
            add_to_vector_db(
                content=file_content,
                meta={
                    "id": doc_id,
                    "name": file.filename,
                    "type": "document",
                    "upload_date": document_data["uploadDate"],
                    "lob": lob_info
                },
                collection_name="documents"
            )
            print(f"✅ [VectorDB] Added to vector DB")

        except Exception as db_error:
            print(f"❌ [DB] Failed to save document: {db_error}")
            import traceback
            traceback.print_exc()
            return jsonify({
                "status": "error",
                "message": f"Failed to save document to the database: {str(db_error)}"
            }), 500
        finally:
            try:
                db.close()
            except Exception:
                pass
        
        print(f"DEBUG: Document upload completed successfully for {file.filename}")
        return jsonify({
            "success": True,
            "message": "Document uploaded successfully",
            # Shape aligned with frontend DocumentsSection expectations
            "id": doc_id,
            "name": file.filename,
            "uploadDate": document_data.get("uploadDate"),
            "fileType": document_data.get("fileType"),
            "status": document_data.get("status"),
            # Additional fields
            "file_path": file_path,
            "size": file_size,
            "content_length": len(file_content) if file_content else 0,
            "lob": lob_info
        }), 201
        
    except Exception as e:
        print(f"❌ Error uploading document: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed to upload document: {str(e)}"}), 500

        # Read file content once and store it
        file_content_bytes = file.read()
        file_size = len(file_content_bytes)
        print(f"DEBUG: File size: {file_size} bytes")
        
        # Create uploads directory if it doesn't exist
        uploads_dir = "uploads"
        if not os.path.exists(uploads_dir):
            os.makedirs(uploads_dir)
            print(f"📁 Created uploads directory: {uploads_dir}")
        
        # Extract content from file
        file_content, images, error = agent_extract_content(io.BytesIO(file_content_bytes), file.filename)
        if error:
            print(f"DEBUG: Content extraction failed: {error}")
            return jsonify({"error": f"Failed to extract content: {error}"}), 500

        print(f"DEBUG: Content extracted successfully, length: {len(file_content)} characters")

        # Create document record
        doc_id = str(uuid.uuid4())
        # Now persist original upload to disk under generated ID
        file_path = f"{uploads_dir}/{doc_id}_{file.filename}"
        with open(file_path, 'wb') as f:
            f.write(file_content_bytes)
        print(f"💾 [File] Saved file to disk: {file_path}")
        # Classify LOB for the document
        lob_info = classify_line_of_business(file_content)

        document_data = {
            "id": doc_id,
            "name": file.filename,
            "uploadDate": datetime.now().isoformat(),
            "fileType": file.filename.split('.')[-1].lower(),
            "size": file_size,
            "status": "uploaded",
            "lob": lob_info
        }
        
        print(f"DEBUG: Document data created with ID: {doc_id}")
        
        # Save to database using SQLAlchemy, add to vector DB with LOB metadata
        db = next(get_db())
        try:
            print(f"💾 [DB] Saving document to database using SQLAlchemy: {file.filename}")
            save_document_to_db(db, document_data, file_path, file_content)
            print(f"✅ [DB] Document stored with ID: {doc_id}")

            # Add to vector database
            print(f"📡 [VectorDB] Adding to vector DB...")
            add_to_vector_db(
                content=file_content,
                meta={
                    "id": doc_id,
                    "name": file.filename,
                    "type": "document",
                    "upload_date": document_data["uploadDate"],
                    "lob": lob_info
                },
                collection_name="documents"
            )
            print(f"✅ [VectorDB] Added to vector DB")

        except Exception as db_error:
            print(f"❌ [DB] Failed to save document: {db_error}")
            import traceback
            traceback.print_exc()
            raise db_error
        finally:
            try:
                db.close()
            except Exception:
                pass


        
        print(f"DEBUG: Document upload completed successfully for {file.filename}")
        return jsonify({
            "success": True,
            "message": "Document uploaded successfully",
            # Shape aligned with frontend DocumentsSection expectations
            "id": doc_id,
            "name": file.filename,
            "uploadDate": document_data.get("uploadDate"),
            "fileType": document_data.get("fileType"),
            "status": document_data.get("status"),
            # Additional fields
            "file_path": file_path,
            "size": file_size,
            "content_length": len(file_content) if file_content else 0,
            "lob": lob_info
        }), 201
        
    except Exception as e:
        print(f"❌ Error uploading document: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed to upload document: {str(e)}"}), 500

@app.route("/api/test_db", methods=['GET'])
def test_database():
    """Test database connectivity and document count"""
    try:
        db = next(get_db())
        try:
            # Test document count
            doc_count = db.query(Document).count()
            analysis_count = db.query(Analysis).count()
            
            return jsonify({
                "success": True,
                "message": "Database connection successful",
                "document_count": doc_count,
                "analysis_count": analysis_count,
                "database_url": DATABASE_URL.split('@')[0] + "@[HIDDEN]"  # Hide credentials
            })
        finally:
            db.close()
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"Database connection failed: {str(e)}",
            "error": str(e)
        }), 500

@app.route("/api/test_create_doc", methods=['POST'])
def test_create_document():
    """Test endpoint to manually create a document in the database"""
    try:
        data = request.get_json()
        test_content = data.get('content', 'Test document content')
        test_filename = data.get('filename', 'test_document.txt')
        
        doc_id = str(uuid.uuid4())
        document_data = {
            "id": doc_id,
            "name": test_filename,
            "uploadDate": datetime.now().isoformat(),
            "fileType": "txt",
            "size": len(test_content),
            "status": "uploaded"
        }
        
        print(f"DEBUG: Testing document creation with ID: {doc_id}")
        
        db = next(get_db())
        try:
            save_document_to_db(db, document_data, f"test_uploads/{doc_id}_{test_filename}", test_content)
            
            return jsonify({
                "success": True,
                "message": "Test document created successfully",
                "document_id": doc_id,
                "document_data": document_data
            })
        finally:
            db.close()
            
    except Exception as e:
        print(f"Error in test document creation: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed to create test document: {str(e)}"}), 500

@app.route("/api/documents", methods=['GET'])
def get_documents():
    """Get all documents from database using direct psycopg2"""
    try:
        documents = get_all_documents_from_db_direct()
        print(f"DEBUG: Retrieved {len(documents)} documents from database using direct psycopg2")
        return jsonify(documents)
    except Exception as e:
        print(f"❌ Error retrieving documents: {e}")
        return jsonify({"error": f"Failed to retrieve documents: {str(e)}"}), 500

@app.route("/api/documents/files", methods=['GET'])
def list_uploaded_files():
    """List all files in the uploads directory"""
    try:
        uploads_dir = "uploads"
        if not os.path.exists(uploads_dir):
            return jsonify({"files": [], "message": "Uploads directory does not exist"})
        
        files = []
        for filename in os.listdir(uploads_dir):
            file_path = os.path.join(uploads_dir, filename)
            if os.path.isfile(file_path):
                file_size = os.path.getsize(file_path)
                files.append({
                    "name": filename,
                    "size": file_size,
                    "path": file_path,
                    "modified": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
                })
        
        return jsonify({
            "files": files,
            "count": len(files),
            "directory": uploads_dir
        })
    except Exception as e:
        return jsonify({"error": f"Failed to list files: {str(e)}"}), 500

@app.route("/api/documents/<doc_id>", methods=['GET'])
def get_document(doc_id):
    """Get a specific document by ID"""
    try:
        db = next(get_db())
        try:
            doc = get_document_by_id(db, doc_id)
            if doc:
                return jsonify(doc)
            else:
                return jsonify({"error": "Document not found"}), 404
        finally:
            db.close()
    except Exception as e:
        return jsonify({"error": f"Failed to retrieve document: {str(e)}"}), 500

@app.route("/api/documents/<doc_id>/file", methods=['GET'])
def download_document(doc_id):
    """Download the actual file for a document"""
    try:
        db = next(get_db())
        try:
            doc = get_document_by_id(db, doc_id)
            if doc and os.path.exists(doc['file_path']):
                return send_file(doc['file_path'], as_attachment=True, download_name=doc['name'])
            else:
                return jsonify({"error": "Document file not found"}), 404
        finally:
            db.close()
    except Exception as e:
        return jsonify({"error": f"Failed to download document: {str(e)}"}), 500

# --- Past Analyses Management ---
@app.route("/api/analyses", methods=['GET'])
def get_past_analyses():
    """Get past analyses from database with pagination"""
    try:
        db = next(get_db())
        try:
            # Get pagination parameters from query string
            limit = int(request.args.get('limit', 50))
            offset = int(request.args.get('offset', 0))
            
            analyses = get_all_analyses_from_db(db, limit=limit, offset=offset)
            print(f"DEBUG: Retrieved {len(analyses)} analyses from database (limit={limit}, offset={offset})")
            return jsonify(analyses)
        finally:
            db.close()
    except Exception as e:
        print(f"❌ Error retrieving analyses: {e}")
        return jsonify({"error": f"Failed to retrieve analyses: {str(e)}"}), 500

@app.route("/api/analyses/<analysis_id>", methods=['GET'])
def get_analysis(analysis_id):
    """Get a specific analysis by ID from database"""
    try:
        db = next(get_db())
        try:
            analysis = get_analysis_by_id_from_db(db, analysis_id)
            if analysis:
                print(f"DEBUG: Retrieved analysis '{analysis['title']}' from database")
                return jsonify(analysis)
            else:
                print(f"DEBUG: Analysis with ID '{analysis_id}' not found")
                return jsonify({"error": "Analysis not found"}), 404
        finally:
            db.close()
    except Exception as e:
        print(f"❌ Error retrieving analysis: {e}")
        return jsonify({"error": f"Failed to retrieve analysis: {str(e)}"}), 500

@app.route("/api/analyses/<analysis_id>/details", methods=['GET'])
def get_analysis_details(analysis_id):
    """Get full analysis details including original_text and results"""
    try:
        db = next(get_db())
        try:
            analysis = get_analysis_details_from_db(db, analysis_id)
            if analysis:
                print(f"DEBUG: Retrieved full analysis details for '{analysis['title']}'")
                return jsonify(analysis)
            else:
                print(f"DEBUG: Analysis with ID '{analysis_id}' not found")
                return jsonify({"error": "Analysis not found"}), 404
        finally:
            db.close()
    except Exception as e:
        print(f"❌ Error retrieving analysis details: {e}")
        return jsonify({"error": f"Failed to retrieve analysis details: {str(e)}"}), 500

def get_analysis_by_id(analysis_id):
    """Get analysis by ID from database"""
    try:
        db = next(get_db())
        try:
            analysis = get_analysis_by_id_from_db(db, analysis_id)
            return analysis
        finally:
            db.close()
    except Exception as e:
        print(f"❌ Error getting analysis by ID: {e}")
        return None

def save_analysis_results(results, original_text, filename, user_email=None, document_id=None):
    """Save analysis results to database"""
    try:
        analysis_id = str(uuid.uuid4())
        analysis_data = {
            "id": analysis_id,
            "title": f"Analysis of {filename}",
            "originalText": original_text,
            "results": results,
            "user_email": user_email,
            "document_id": document_id
        }
        
        db = next(get_db())
        try:
            save_analysis_to_db(db, analysis_data)
            print(f"✅ Analysis saved to database: {analysis_id}")
            return analysis_id
        finally:
            db.close()
            
    except Exception as e:
        print(f"❌ Error saving analysis results: {e}")
        return None

# --- Helper function to convert markdown to DOCX ---
def markdown_to_docx(markdown_content):
    """Convert markdown content to DOCX format with enhanced table support"""
    doc = DocxDocument()
    
    # Split content into lines
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            doc.add_paragraph()
            i += 1
            continue
            
        # Handle headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if level == 1:
                doc.add_heading(text, 0)
            elif level == 2:
                doc.add_heading(text, 1)
            elif level == 3:
                doc.add_heading(text, 2)
            elif level == 4:
                doc.add_heading(text, 3)
            else:
                doc.add_heading(text, 4)
            i += 1
            
        # Handle tables
        elif line.startswith('|'):
            # Start of a table
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) >= 2:  # Need at least header and separator
                # Parse table
                table_data = []
                for table_line in table_lines:
                    # Handle both |text| and text|text formats
                    if table_line.startswith('|'):
                        # Remove leading pipe if present
                        clean_line = table_line[1:] if table_line.startswith('|') else table_line
                        # Remove trailing pipe if present
                        clean_line = clean_line[:-1] if clean_line.endswith('|') else clean_line
                        # Split by pipes and clean up each cell
                        cells = clean_line.split('|')
                        cells = [cell.strip() for cell in cells]
                        table_data.append(cells)
                
                if table_data:
                    # Create table
                    num_rows = len(table_data)
                    num_cols = max(len(row) for row in table_data) if table_data else 0
                    
                    if num_rows > 0 and num_cols > 0:
                        table = doc.add_table(rows=num_rows, cols=num_cols)
                        table.style = 'Table Grid'
                        
                        # Fill table data
                        for row_idx, row_data in enumerate(table_data):
                            for col_idx, cell_text in enumerate(row_data):
                                if col_idx < num_cols:
                                    cell = table.cell(row_idx, col_idx)
                                    cell.text = cell_text
                                    
                                    # Make header row bold
                                    if row_idx == 0:
                                        for paragraph in cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.bold = True
                                    
                                    # Apply cell formatting for better readability
                                    for paragraph in cell.paragraphs:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                        for run in paragraph.runs:
                                            run.font.size = Pt(10)
        
        # Handle bold text
        elif '**' in line:
            # Simple bold handling - replace **text** with bold
            p = doc.add_paragraph()
            parts = line.split('**')
            for j, part in enumerate(parts):
                if j % 2 == 1:  # Bold text
                    p.add_run(part).bold = True
                else:  # Regular text
                    p.add_run(part)
            i += 1
            
        # Handle lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        elif line.startswith('1. '):
            doc.add_paragraph(line[3:], style='List Number')
            i += 1
            
        # Handle code blocks
        elif line.startswith('```'):
            # Skip code block markers
            i += 1
            # Add code block content as preformatted text
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_line = lines[i]
                p = doc.add_paragraph(code_line)
                # Apply monospace font to code
                for run in p.runs:
                    run.font.name = 'Courier New'
                i += 1
            if i < len(lines) and lines[i].strip().startswith('```'):
                i += 1  # Skip closing ```
                
        # Handle regular paragraphs
        else:
            doc.add_paragraph(line)
            i += 1
    
    return doc

def convert_image_format(image_data, from_mime_type, to_mime_type='image/png'):
    """Convert image from one format to another (enhanced implementation)"""
    try:
        # Try to import PIL for image conversion
        try:
            from PIL import Image, ImageDraw
            import io
            
            # For WMF files, we need special handling
            if from_mime_type.lower() in ['image/x-wmf', 'image/wmf']:
                print(f"DEBUG: WMF format detected - attempting enhanced handling...")
                
                # Method 1: Try direct PIL opening (rarely works for WMF)
                try:
                    image = Image.open(io.BytesIO(image_data))
                    print(f"DEBUG: Successfully opened WMF with PIL")
                except Exception as wmf_error:
                    print(f"DEBUG: PIL cannot handle WMF directly: {wmf_error}")
                    
                    # Method 2: Try to install and use wand (ImageMagick wrapper)
                    try:
                        from wand.image import Image as WandImage
                        print(f"DEBUG: Attempting WMF conversion with ImageMagick...")
                        
                        with WandImage(blob=image_data) as wand_img:
                            # Convert to PNG format
                            wand_img.format = 'png'
                            converted_data = wand_img.make_blob()
                            print(f"DEBUG: Successfully converted WMF to PNG using ImageMagick")
                            return converted_data, 'image/png'
                            
                    except ImportError:
                        print(f"DEBUG: ImageMagick (wand) not available for WMF conversion")
                    except Exception as magick_error:
                        print(f"DEBUG: ImageMagick conversion failed: {magick_error}")
                    
                    # Method 3: Try using Windows-specific libraries (if available)
                    try:
                        import win32com.client
                        print(f"DEBUG: Attempting WMF conversion using Windows COM...")
                        # This would require Windows-specific code
                        print(f"DEBUG: Windows COM conversion not implemented yet")
                    except ImportError:
                        print(f"DEBUG: Windows COM libraries not available")
                    except Exception as com_error:
                        print(f"DEBUG: Windows COM conversion failed: {com_error}")
                    
                    # Method 4: Create a placeholder image
                    print(f"DEBUG: Creating placeholder image for WMF...")
                    try:
                        # Create a simple placeholder image
                        placeholder = Image.new('RGB', (300, 200), color='lightgray')
                        draw = ImageDraw.Draw(placeholder)
                        draw.text((10, 10), "WMF Image", fill='black')
                        draw.text((10, 30), "Format not supported", fill='red')
                        
                        output_buffer = io.BytesIO()
                        placeholder.save(output_buffer, format='PNG')
                        converted_data = output_buffer.getvalue()
                        print(f"DEBUG: Created placeholder image for WMF")
                        return converted_data, 'image/png'
                        
                    except Exception as placeholder_error:
                        print(f"DEBUG: Failed to create placeholder: {placeholder_error}")
                    
                    print(f"DEBUG: All WMF conversion methods failed. Skipping image.")
                    return None, None
            
            # For other formats, use standard PIL conversion
            # Convert bytes to PIL Image
            image = Image.open(io.BytesIO(image_data))
            
            # Convert to RGB if needed
            if image.mode in ('RGBA', 'LA', 'P'):
                image = image.convert('RGB')
            
            # Convert to bytes in the target format
            output_buffer = io.BytesIO()
            if to_mime_type == 'image/png':
                image.save(output_buffer, format='PNG')
            elif to_mime_type == 'image/jpeg':
                image.save(output_buffer, format='JPEG', quality=85)
            else:
                image.save(output_buffer, format='PNG')  # Default to PNG
            
            converted_data = output_buffer.getvalue()
            print(f"DEBUG: Successfully converted {from_mime_type} to {to_mime_type}")
            return converted_data, to_mime_type
            
        except ImportError:
            print("DEBUG: PIL not available, skipping image conversion")
            return None, None
        except Exception as e:
            print(f"DEBUG: Image conversion failed: {e}")
            return None, None
            
    except Exception as e:
        print(f"DEBUG: Error in image conversion: {e}")
        return None, None

def is_valid_image_for_gemini(mime_type, image_data):
    """Check if image is valid for Gemini API and convert if needed"""
    # Gemini supports: image/jpeg, image/png, image/webp, image/heic, image/heif
    supported_types = ['image/jpeg', 'image/png', 'image/webp', 'image/heic', 'image/heif']
    is_supported = mime_type.lower() in supported_types
    
    if not is_supported:
        print(f"DEBUG: Image rejected - unsupported MIME type: {mime_type}")
        # Try to convert unsupported formats
        if mime_type.lower() in ['image/x-wmf', 'image/wmf', 'image/bmp', 'image/tiff']:
            print(f"DEBUG: Attempting to convert {mime_type} to PNG...")
            converted_data, converted_mime = convert_image_format(image_data, mime_type, 'image/png')
            if converted_data:
                print(f"DEBUG: Successfully converted {mime_type} to PNG")
                return True  # Return True since we now have a supported format
            else:
                print(f"DEBUG: Conversion failed for {mime_type}")
                return False
    else:
        print(f"DEBUG: Image accepted - supported MIME type: {mime_type}")
    
    return is_supported

def process_images_for_gemini(images):
    """Process images to ensure they are compatible with Gemini API using multiple methods"""
    print(f"🔧 Processing {len(images)} images with enhanced methods...")
    
    # Use enhanced processing if available
    if sum(ENHANCED_IMAGE_LIBRARIES.values()) > 1:
        print("DEBUG: Using enhanced image processing with multiple libraries")
        return process_images_with_multiple_methods(images)
    else:
        print("DEBUG: Using basic image processing")
        return process_images_for_gemini_basic(images)

def process_images_for_gemini_basic(images):
    """Basic image processing (original method)"""
    valid_images = []
    for i, img in enumerate(images):
        if is_valid_image_for_gemini(img["mime_type"], img["data"]):
            valid_images.append(img)
            print(f"DEBUG: Image {i+1} is valid for Gemini API: {img['mime_type']}")
        else:
            print(f"DEBUG: Image {i+1} skipped - unsupported format: {img['mime_type']}")
    return valid_images

# --- File Extraction Agent ---

def agent_extract_content(file_stream, filename):
    """Extract text content and images from uploaded files with enhanced image processing"""
    print("AGENT [Extractor]: Starting enhanced content extraction...")
    
    try:
        file_extension = filename.lower().split('.')[-1]
        
        if file_extension == 'docx':
            # Handle DOCX files with image extraction
            try:
                doc = DocxDocument(file_stream)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                images = []
                
                # Extract images from DOCX
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_data = rel.target_part.blob
                            mime_type = rel.target_part.content_type
                            
                            if is_valid_image_for_gemini(mime_type, image_data):
                                images.append({
                                    'mime_type': mime_type,
                                    'data': base64.b64encode(image_data).decode('utf-8')
                                })
                                print(f"AGENT [Extractor]: Extracted image with MIME type: {mime_type}")
                        except Exception as e:
                            print(f"AGENT [Extractor]: Failed to extract image: {e}")
                
                print(f"AGENT [Extractor]: Extracted text ({len(text_content)} chars) and {len(images)} images from DOCX.")
                
                # Log token consumption for content extraction
                estimated_tokens = len(text_content.split()) * 1.3
                log_token_consumption("content_extraction", estimated_tokens, "local_processing", {
                    "file_type": "docx",
                    "text_length": len(text_content),
                    "images_count": len(images)
                })
                
                return text_content, images, None
            except PackageNotFoundError:
                return "", [], "Invalid DOCX file format"
            except Exception as e:
                return "", [], f"Error processing DOCX file: {str(e)}"
        
        elif file_extension == 'pdf':
            # Handle PDF files with image extraction
            try:
                pdf_reader = PyPDF2.PdfReader(file_stream)
                text_content = ""
                images = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text_content += page.extract_text() + "\n"
                    
                    # Extract images from PDF (basic implementation)
                    # Note: PyPDF2 has limited image extraction capabilities
                    # For production, consider using pdfplumber or pdf2image
                    try:
                        if '/XObject' in page['/Resources']:
                            xObject = page['/Resources']['/XObject'].get_object()
                            for obj in xObject:
                                if xObject[obj]['/Subtype'] == '/Image':
                                    try:
                                        image_data = xObject[obj].get_data()
                                        # Determine MIME type based on image format
                                        if xObject[obj]['/Filter'] == '/DCTDecode':
                                            mime_type = 'image/jpeg'
                                        elif xObject[obj]['/Filter'] == '/FlateDecode':
                                            mime_type = 'image/png'
                                        else:
                                            mime_type = 'image/jpeg'  # Default
                                        
                                        if is_valid_image_for_gemini(mime_type, image_data):
                                            images.append({
                                                'mime_type': mime_type,
                                                'data': base64.b64encode(image_data).decode('utf-8')
                                            })
                                            print(f"AGENT [Extractor]: Extracted image from PDF page {page_num + 1}")
                                    except Exception as e:
                                        print(f"AGENT [Extractor]: Failed to extract image from PDF: {e}")
                    except Exception as e:
                        print(f"AGENT [Extractor]: Error processing PDF images: {e}")
                
                print(f"AGENT [Extractor]: Extracted text ({len(text_content)} chars) and {len(images)} images from PDF.")
                
                # Log token consumption for content extraction
                estimated_tokens = len(text_content.split()) * 1.3
                log_token_consumption("content_extraction", estimated_tokens, "local_processing", {
                    "file_type": "pdf",
                    "text_length": len(text_content),
                    "images_count": len(images)
                })
                
                return text_content, images, None
            except Exception as e:
                return "", [], f"Error processing PDF file: {str(e)}"
        
        else:
            return "", [], f"Unsupported file type: {file_extension}"
            
    except Exception as e:
        print(f"AGENT [Extractor]: Error during extraction: {e}")
        return "", [], f"Extraction failed: {str(e)}"

# --- Gemini API Agent Caller ---

def call_generative_agent(prompt_parts, is_json=False, stage_name="generative_agent"):
    """Call the generative AI agent with proper error handling, detailed token tracking, and retry logic"""
    import time
    import random
    
    max_retries = 3
    base_delay = 2  # Base delay in seconds
    
    for attempt in range(max_retries):
        try:
            # Extract text from prompt_parts (which are dictionaries)
            text_parts = []
            total_input_chars = 0
            image_count = 0
            
            for part in prompt_parts:
                if isinstance(part, dict):
                    if 'text' in part:
                        text_parts.append(part['text'])
                        total_input_chars += len(part['text'])
                    elif 'inline_data' in part:
                        # This is an image
                        image_count += 1
                        text_parts.append("[IMAGE]")  # Placeholder for image
                        total_input_chars += 1000  # Rough estimate for image
                else:
                    text_parts.append(str(part))
                    total_input_chars += len(str(part))
            
            total_text = " ".join(text_parts)
            estimated_input_tokens = len(total_text.split()) * 1.3  # Rough token estimation
            
            print(f"🔍 {stage_name.upper()}: Starting API call (attempt {attempt + 1}/{max_retries})")
            print(f"   📝 Input characters: {total_input_chars:,}")
            print(f"   📊 Estimated input tokens: {estimated_input_tokens:,.0f}")
            print(f"   🖼️ Images in prompt: {image_count}")
            print(f"   🔑 API Key: {GEMINI_API_KEY[:10]}...")
            print(f"   🌐 API URL: {GEMINI_API_URL}")
            
            headers = {
                "Content-Type": "application/json",
                "x-goog-api-key": GEMINI_API_KEY
            }
            
            # Prepare the request payload
            if is_json:
                payload = {
                    "contents": [{
                        "role": "user",
                        "parts": [{"text": "You are a helpful AI assistant. Please respond with valid JSON only, no additional text."}]
                    }, {
                        "role": "user",
                        "parts": prompt_parts
                    }],
                    "generationConfig": {
                        "temperature": 0.1,
                        "topK": 1,
                        "topP": 0.1,
                        "maxOutputTokens": 8192,
                        "responseMimeType": "application/json"
                    }
                }
            else:
                payload = {
                    "contents": [{
                        "role": "user",
                        "parts": prompt_parts
                    }],
                    "generationConfig": {
                        "temperature": 0.7,
                        "topK": 40,
                        "topP": 0.95,
                        "maxOutputTokens": 8192
                    }
                }
            
            print(f"   📦 Payload structure: {len(payload['contents'])} content items")
            print(f"   📦 First part type: {type(prompt_parts[0]) if prompt_parts else 'None'}")
            
            response = requests.post(GEMINI_API_URL, headers=headers, json=payload, timeout=120)
            
            print(f"   📡 Response status: {response.status_code}")
            print(f"   📡 Response headers: {dict(response.headers)}")
            
            if response.status_code == 200:
                result = response.json()
                print(f"   ✅ API call successful")
                
                # Extract detailed token information
                usage_metadata = result.get('usageMetadata', {})
                input_token_count = usage_metadata.get('promptTokenCount', estimated_input_tokens)
                output_token_count = usage_metadata.get('candidatesTokenCount', 0)
                total_token_count = usage_metadata.get('totalTokenCount', input_token_count + output_token_count)
                
                # Get response text for output token estimation if not provided
                if output_token_count == 0:
                    content = result.get("candidates", [{}])[0].get("content", {})
                    parts = content.get("parts", [])
                    response_text = parts[0].get("text", "") if parts else ""
                    output_token_count = len(response_text.split()) * 1.3
                
                # Log detailed token consumption
                details = {
                    "input_tokens": input_token_count,
                    "output_tokens": output_token_count,
                    "input_characters": total_input_chars,
                    "response_length": len(response_text) if 'response_text' in locals() else 0,
                    "is_json": is_json,
                    "attempt": attempt + 1,
                    "images_in_prompt": image_count
                }
                
                log_token_consumption(stage_name, total_token_count, "gemini-1.5-pro", details)
                
                if is_json:
                    # Extract JSON from the response
                    content = result.get("candidates", [{}])[0].get("content", {})
                    parts = content.get("parts", [])
                    if parts:
                        text_content = parts[0].get("text", "")
                        print(f"AGENT [{stage_name}]: Raw response length: {len(text_content)}")
                        print(f"AGENT [{stage_name}]: Raw response preview: {text_content[:200]}...")
                        
                        # Try to extract JSON from the response
                        try:
                            # Find JSON in the response - look for the complete JSON structure
                            json_start = text_content.find('{')
                            if json_start != -1:
                                # Find the matching closing brace by counting braces
                                brace_count = 0
                                json_end = json_start
                                for i in range(json_start, len(text_content)):
                                    if text_content[i] == '{':
                                        brace_count += 1
                                    elif text_content[i] == '}':
                                        brace_count -= 1
                                        if brace_count == 0:
                                            json_end = i + 1
                                            break
                                
                                if json_end > json_start:
                                    json_str = text_content[json_start:json_end]
                                    print(f"AGENT [{stage_name}]: Extracted JSON length: {len(json_str)}")
                                    parsed_json = json.loads(json_str)
                                    print(f"AGENT [{stage_name}]: JSON parsing successful")
                                    return parsed_json
                                else:
                                    print(f"AGENT [{stage_name}]: Could not find complete JSON structure")
                                    return None
                            else:
                                print(f"AGENT [{stage_name}]: No JSON structure found in response")
                                return None
                        except json.JSONDecodeError as e:
                            print(f"AGENT [{stage_name}]: JSON parsing failed: {e}")
                            print(f"AGENT [{stage_name}]: Attempted to parse: {text_content[json_start:json_start+500]}...")
                            return None
                    else:
                        print(f"AGENT [{stage_name}]: No content parts found in response")
                        return None
                else:
                    content = result.get("candidates", [{}])[0].get("content", {})
                    parts = content.get("parts", [])
                    return parts[0].get("text", "") if parts else ""
                    
            elif response.status_code == 429:
                # Rate limit exceeded - implement exponential backoff
                error_data = response.json()
                retry_delay = error_data.get('error', {}).get('details', [{}])[0].get('retryDelay', '40s')
                
                # Parse retry delay (e.g., "40s" -> 40 seconds)
                if isinstance(retry_delay, str) and retry_delay.endswith('s'):
                    try:
                        retry_seconds = int(retry_delay[:-1])
                    except ValueError:
                        retry_seconds = base_delay * (2 ** attempt)
                else:
                    retry_seconds = base_delay * (2 ** attempt)
                
                # Add jitter to prevent thundering herd
                jitter = random.uniform(0.5, 1.5)
                actual_delay = retry_seconds * jitter
                
                print(f"⚠️ Rate limit exceeded. Retrying in {actual_delay:.1f} seconds...")
                log_token_consumption(stage_name, 0, "gemini-1.5-pro", {
                    "error": f"Rate limit exceeded (attempt {attempt + 1})",
                    "retry_delay": actual_delay
                })
                
                if attempt < max_retries - 1:
                    time.sleep(actual_delay)
                    continue
                else:
                    print(f"❌ Max retries exceeded for {stage_name}")
                    log_token_consumption(stage_name, 0, "gemini-1.5-pro", {"error": "Max retries exceeded"})
                    return None
                    
            elif response.status_code == 503:
                # Service unavailable - implement longer backoff
                print(f"⚠️ Service unavailable (503). Retrying in {base_delay * (2 ** attempt):.1f} seconds...")
                log_token_consumption(stage_name, 0, "gemini-1.5-pro", {
                    "error": f"Service unavailable (503) (attempt {attempt + 1})",
                    "retry_delay": base_delay * (2 ** attempt)
                })
                
                if attempt < max_retries - 1:
                    time.sleep(base_delay * (2 ** attempt))
                    continue
                else:
                    print(f"❌ Max retries exceeded for {stage_name} (503 errors)")
                    log_token_consumption(stage_name, 0, "gemini-1.5-pro", {"error": "Max retries exceeded (503)"})
                    return None
                    
            else:
                print(f"❌ API Error: {response.status_code} - {response.text}")
                log_token_consumption(stage_name, 0, "gemini-1.5-pro", {"error": f"API Error {response.status_code}"})
                return None
                
        except Exception as e:
            print(f"❌ Error calling generative agent (attempt {attempt + 1}): {e}")
            import traceback
            traceback.print_exc()
            log_token_consumption(stage_name, 0, "gemini-1.5-pro", {"error": str(e), "attempt": attempt + 1})
            
            if attempt < max_retries - 1:
                delay = base_delay * (2 ** attempt) + random.uniform(0, 1)
                print(f"⚠️ Retrying in {delay:.1f} seconds...")
                time.sleep(delay)
                continue
            else:
                print(f"❌ Max retries exceeded for {stage_name}")
                return None
    
    return None

# --- Specialized Generative Agents ---

def agent_planner(text_content, images):
    print("AGENT [Planner]: Starting enhanced analysis to create a comprehensive plan...")
    prompt_text = f"""You are an expert Business Analyst with 15+ years of experience in enterprise software development. 
    
Analyze the following business requirements document thoroughly, including its text and any images. Create a comprehensive, detailed plan that covers:

1. **System Overview**: High-level architecture and main components
2. **User Roles & Permissions**: All user types and their access levels
3. **Core Features**: Primary functionality and user workflows
4. **Technical Requirements**: Non-functional requirements (performance, security, scalability)
5. **Integration Points**: External systems and APIs
6. **Data Requirements**: Data models, storage, and migration needs
7. **Security Considerations**: Authentication, authorization, and compliance
8. **Deployment Strategy**: Infrastructure and deployment approach

Focus on creating a plan that is:
- Comprehensive yet concise
- Actionable for development teams
- Aligned with industry best practices
- Scalable for future enhancements

--- DOCUMENT TEXT ---
{text_content}
--- DOCUMENT IMAGES ---
"""
    prompt_parts = [{"text": prompt_text}]
    
    # Process images to ensure they are compatible with Gemini API
    valid_images = process_images_for_gemini(images)
    
    if valid_images:
        print(f"AGENT [Planner]: Including {len(valid_images)} valid images in the analysis")
        for img in valid_images:
            prompt_parts.append({
                "inline_data": {
                    "mime_type": img["mime_type"],
                    "data": img["data"]
                }
            })
    else:
        print("AGENT [Planner]: No valid images found, proceeding with text-only analysis")
        # Update prompt to indicate no images were included
        prompt_parts[0]["text"] = prompt_text.replace("including its text and any images", "including its text (no valid images were found)")
    
    plan = call_generative_agent(prompt_parts, stage_name="planning_analysis")
    if plan is None:
        # Fallback: Create a basic plan from the text content
        print("AGENT [Planner]: API failed, creating fallback plan...")
        fallback_plan = f"""# FALLBACK ANALYSIS PLAN

## System Overview
Based on the provided requirements document, this system appears to be a business application requiring analysis.

## Key Components Identified
- User interface requirements
- Data processing needs
- System integration requirements

## Technical Considerations
- Performance requirements
- Security implementation
- Scalability planning

## Next Steps
- Detailed requirements analysis needed
- Technical architecture design required
- Implementation planning necessary

Note: This is a fallback plan due to API unavailability. For comprehensive analysis, please try again when the service is available.

--- ORIGINAL TEXT ---
{text_content[:1000]}...
"""
        error = None
        print("AGENT [Planner]: Fallback plan created successfully.")
        return fallback_plan, error
    else:
        error = None
        print("AGENT [Planner]: Enhanced plan created successfully.")
        return plan, error

def agent_trd_writer(plan, original_text, lob_info=None):
    """Enhanced TRD writer with better structure and content, focusing on Non-Functional Requirements instead of LOB"""
    print("AGENT [TRD Writer]: Creating comprehensive Technical Requirements Document...")
    
    # Enhanced Non-Functional Requirements section
    nfr_section = """
## 4. NON-FUNCTIONAL REQUIREMENTS

### 4.1 Performance Requirements
- Response Time: System must respond within 2 seconds for 95% of requests
- Throughput: Support 1000+ concurrent users
- Scalability: Handle 10x growth in data volume
- Availability: 99.9% uptime requirement

### 4.2 Security Requirements
- Authentication: Multi-factor authentication (MFA) support
- Authorization: Role-based access control (RBAC)
- Data Encryption: AES-256 encryption at rest and in transit
- Audit Logging: Complete audit trail for all user actions
- Compliance: GDPR, SOC2, and industry-specific regulations

### 4.3 Scalability Requirements
- Horizontal scaling capability
- Load balancing support
- Database sharding and partitioning
- Microservices architecture support
- Cloud-native deployment ready

### 4.4 Usability Requirements
- User Interface: Intuitive and responsive design
- Accessibility: WCAG 2.1 AA compliance
- Mobile Support: Responsive design for all devices
- Internationalization: Multi-language support
- User Experience: 90% user satisfaction target

### 4.5 Reliability Requirements
- Fault tolerance and error handling
- Graceful degradation under load
- Backup and disaster recovery
- Data integrity and consistency
- Service level agreements (SLAs)

### 4.6 Maintainability Requirements
- Code quality and documentation standards
- Automated testing and CI/CD pipelines
- Monitoring and alerting systems
- Performance metrics and dashboards
- Technical debt management
"""

    prompt = f"""You are a Senior Technical Writer and Business Analyst with 20+ years of experience in enterprise software documentation.

Create a comprehensive Technical Requirements Document (TRD) based on the following plan and original requirements. Structure it as follows:

# TECHNICAL REQUIREMENTS DOCUMENT

## 1. EXECUTIVE SUMMARY
- Project Overview
- Key Objectives
- Success Criteria
- Business Value Proposition

## 2. SYSTEM OVERVIEW
- High-Level Architecture
- System Components
- Technology Stack
- Integration Points

## 3. FUNCTIONAL REQUIREMENTS
- User Stories and Use Cases
- Business Rules
- Data Requirements
- Workflow Definitions
- Business Process Mapping

{nfr_section}

## 5. TECHNICAL SPECIFICATIONS
- API Specifications (REST/GraphQL)
- Database Design and Schema
- Integration Points and APIs
- Third-party Service Dependencies
- Data Flow and Processing

## 6. SECURITY CONSIDERATIONS
- Authentication & Authorization
- Data Protection and Privacy
- Compliance Requirements
- Security Testing and Penetration Testing
- Incident Response Procedures

## 7. DEPLOYMENT & OPERATIONS
- Infrastructure Requirements
- Deployment Strategy (DevOps/DevSecOps)
- Monitoring & Maintenance
- Performance Optimization
- Disaster Recovery Plan

## 8. TESTING STRATEGY
- Test Types (Unit, Integration, E2E)
- Quality Assurance Processes
- Acceptance Criteria
- Performance Testing
- Security Testing

## 9. RISK ASSESSMENT
- Technical Risks
- Business Risks
- Mitigation Strategies
- Contingency Plans

## 10. SUCCESS METRICS
- Key Performance Indicators (KPIs)
- Business Metrics
- Technical Metrics
- User Satisfaction Metrics

Make the document:
- Comprehensive yet concise
- Technical but accessible
- Actionable for development teams
- Compliant with industry standards
- Focus on measurable requirements

--- HIGH-LEVEL PLAN ---
{plan}

--- ORIGINAL REQUIREMENTS TEXT ---
{original_text}
"""
    result = call_generative_agent([{"text": prompt}], stage_name="trd_writer")
    if result is None:
        return "TRD could not be completed due to API error.", "API call failed"
    return result, None

def agent_diagrammer(plan, diagram_type):
    """Generate diagrams in Mermaid; on failure, return a minimal valid diagram."""
    print(f"AGENT [Diagrammer]: Creating {diagram_type} diagram...")
    
    try:

        if diagram_type == "HLD":
            prompt = f"""You are a Senior Solution Architect specializing in system design diagrams.

Create a High-Level Design (HLD) diagram using Mermaid syntax. The diagram should show:

**System Architecture Components:**
- User Interface Layer (Web/Mobile/Desktop)
- Application Layer (API Gateway, Services)
- Business Logic Layer (Core Services)
- Data Access Layer (Repositories)
- External Systems Integration
- Security Layer (Authentication/Authorization)

**Key Elements:**
- System boundaries and data flow
- Integration points with external systems
- Technology stack components
- Security layers and authentication
- Load balancers and caching

Use Mermaid flowchart syntax with:
- Clear component names in boxes
- Proper relationships with arrows
- Color coding for different layers using classDef and class assignments
- Professional styling and layout with consistent colors for layers
- AVOID subgraph syntax - use simple flowchart instead

IMPORTANT: 
- Respond ONLY with the Mermaid diagram code, no additional text or explanations
- Use simple flowchart syntax, NOT subgraph
- Use clear, simple node labels
- Avoid complex syntax that might cause parsing errors

Use THIS CANONICAL MERMAID STYLE as reference. Reuse its init block, classDefs, and linkStyle conventions; only change node labels and edges for the current system. Keep linkStyle indices contiguous starting at 0.

                    ```mermaid
                    %%{init: {
                      "theme": "default",
                      "themeVariables": {
                        "fontFamily": "Segoe UI, Arial",
                        "primaryColor": "#f2f2f2",
                        "edgeLabelBackground":"#f7ffe2"
                      }
                    }}%%
                    flowchart TD
                        start((Start)):::startStyle
                        step1["First Step"]:::stepA
                        dec1{{Decision Point}}:::decision
                        step2["Second Step"]:::stepB
                        endcircle(("END")):::endStyle

                        start --> step1
                        step1 -- Yes --> dec1
                        step1 -- No --> endcircle
                        dec1 -- True --> step2
                        dec1 -- False --> endcircle
                        step2 --> endcircle

                        %% Node styling
                        classDef startStyle fill:#6be585,stroke:#247346,stroke-width:3px,color:#fff;
                        classDef decision fill:#ffd966,stroke:#b29928,stroke-width:3px,color:#593E00;
                        classDef stepA fill:#bbdefb,stroke:#1976d2,stroke-width:2px;
                        classDef stepB fill:#c5e1a5,stroke:#558b2f,stroke-width:2px;
                        classDef endStyle fill:#e57373,stroke:#b71c1c,stroke-width:4px,color:#fff;

                        %% Link styling (by order)
                        linkStyle 0 stroke:#29a19c,stroke-width:4px;
                        linkStyle 1 stroke:#fbc02d,stroke-width:3px,stroke-dasharray: 4 2;
                        linkStyle 2 stroke:#b71c1c,stroke-width:3px,stroke-dasharray: 2 4;
                        linkStyle 3 stroke:#388e3c,stroke-width:3px;
                        linkStyle 4 stroke:#b71c1c,stroke-width:3px,stroke-dasharray: 4 2;
                    ```

--- HIGH-LEVEL PLAN ---
{plan}
"""
        else:  # LLD
            prompt = f"""You are a Senior Software Architect specializing in detailed system design.

Create a Low-Level Design (LLD) diagram using Mermaid syntax. The diagram should show:

**Detailed System Components:**
- Specific modules, classes, and interfaces
- Database schema relationships and tables
- API endpoints and HTTP methods
- Data flow between components
- Error handling and validation paths
- Caching and performance optimizations

**Key Elements:**
- Detailed component interactions
- Database tables and relationships
- API specifications with methods
- Security implementation details
- Performance considerations
- Error handling flows

Use Mermaid flowchart syntax with:
- Detailed component breakdown
- Clear data flow with labels
- Error handling paths
- Performance indicators
- Database relationships
- Color-coded components using classDef and class assignments
- AVOID subgraph syntax - use simple flowchart instead

IMPORTANT: 
- Respond ONLY with the Mermaid diagram code, no additional text or explanations
- Use simple flowchart syntax, NOT subgraph
- Use clear, simple node labels
- Avoid complex syntax that might cause parsing errors

Use THIS CANONICAL MERMAID STYLE as reference. Reuse its init block, classDefs, and linkStyle conventions; only change node labels and edges for the current feature. Keep linkStyle indices contiguous starting at 0.

                    ```mermaid
                    %%{init: {
                      "theme": "default",
                      "themeVariables": {
                        "fontFamily": "Segoe UI, Arial",
                        "primaryColor": "#f2f2f2",
                        "edgeLabelBackground":"#f7ffe2"
                      }
                    }}%%
                    flowchart TD
                        start((Start)):::startStyle
                        step1["First Step"]:::stepA
                        dec1{{Decision Point}}:::decision
                        step2["Second Step"]:::stepB
                        endcircle(("END")):::endStyle

                        start --> step1
                        step1 -- Yes --> dec1
                        step1 -- No --> endcircle
                        dec1 -- True --> step2
                        dec1 -- False --> endcircle
                        step2 --> endcircle

                        %% Node styling
                        classDef startStyle fill:#6be585,stroke:#247346,stroke-width:3px,color:#fff;
                        classDef decision fill:#ffd966,stroke:#b29928,stroke-width:3px,color:#593E00;
                        classDef stepA fill:#bbdefb,stroke:#1976d2,stroke-width:2px;
                        classDef stepB fill:#c5e1a5,stroke:#558b2f,stroke-width:2px;
                        classDef endStyle fill:#e57373,stroke:#b71c1c,stroke-width:4px,color:#fff;

                        %% Link styling (by order)
                        linkStyle 0 stroke:#29a19c,stroke-width:4px;
                        linkStyle 1 stroke:#fbc02d,stroke-width:3px,stroke-dasharray: 4 2;
                        linkStyle 2 stroke:#b71c1c,stroke-width:3px,stroke-dasharray: 2 4;
                        linkStyle 3 stroke:#388e3c,stroke-width:3px;
                        linkStyle 4 stroke:#b71c1c,stroke-width:3px,stroke-dasharray: 4 2;
                    ```

--- HIGH-LEVEL PLAN ---
{plan}
"""
        
        print(f"AGENT [Diagrammer]: Sending {diagram_type} diagram request...")
        result = call_generative_agent([{"text": prompt}], stage_name=f"{diagram_type.lower()}_diagram")
        
        if result is None:
            print(f"AGENT [Diagrammer]: {diagram_type} diagram generation failed - API returned None")
            return f"Failed to generate {diagram_type} diagram due to API error.", "API call failed"
        
        # Clean the result to extract Mermaid code
        cleaned_result = extract_mermaid_code(result)

        # If the model did not return usable code, provide a minimal fallback
        import re
        has_flowchart = False
        if cleaned_result:
            # Accept diagrams that include an init block before the flowchart
            # e.g., starting with `%%{init: ...}%%` followed by `flowchart TD`
            has_flowchart = bool(re.search(r'(^|\n)flowchart\s+\w+', cleaned_result))
        if not cleaned_result or not has_flowchart:
            print(f"AGENT [Diagrammer]: Falling back to contextual minimal {diagram_type} diagram")
            minimal = build_contextual_minimal_diagram(plan or "", diagram_type)
            return minimal, None

        print(f"AGENT [Diagrammer]: Successfully generated {diagram_type} diagram")
        return cleaned_result, None
        
    except Exception as e:
        print(f"AGENT [Diagrammer]: Exception during {diagram_type} generation: {e}")
        import traceback
        traceback.print_exc()
        # Contextual minimal fallback on exception
        minimal = build_contextual_minimal_diagram(plan or "", diagram_type)
        return minimal, None

def build_contextual_minimal_diagram(plan_text: str, diagram_type: str) -> str:
    """Build a small but context-aware Mermaid flowchart using keywords from plan.

    Uses descriptive IDs and simple classDefs to keep rendering stable.
    """
    p = (plan_text or "").lower()

    def has(*keys: str) -> bool:
        return any(k in p for k in keys)

    # Common detections
    ui = has("ui", "user interface", "frontend", "web")
    api = has("api", "gateway", "endpoint")
    auth = has("auth", "authentication", "oauth", "jwt")
    core = has("core", "service", "business logic", "orchestrator")
    repo = has("repo", "repository", "dao", "data access")
    db = has("db", "database", "table", "schema")
    cache = has("cache", "caching")
    lb = has("load balancer", "load-balancer", "lb")

    # External systems
    ext_ping = has("ping")
    ext_touchstone = has("touchstone")
    ext_guidewire = has("guidewire")
    ext_imageright = has("imageright")
    ext_genai = has("genai", "gen ai", "llm", "ai platform")

    lines = ["flowchart TD"]

    if diagram_type.upper() == "HLD":
        # Nodes
        if ui:
            lines.append("UI_UserInterface[User Interface]")
        if api:
            lines.append("APP_ApiGateway[API Gateway]")
        if auth:
            lines.append("SEC_Auth[Authentication]")
        if core:
            lines.append("BIZ_CoreServices[Core Services]")
        if repo:
            lines.append("DATA_Repositories[Repositories]")
        if db:
            lines.append("DATA_Database[(Database)]")
        if lb:
            lines.append("INFRA_LoadBalancer[Load Balancer]")
        if cache:
            lines.append("INFRA_Cache[Cache]")
        if ext_ping:
            lines.append("EXT_PING[Ping]")
        if ext_touchstone:
            lines.append("EXT_Touchstone[Touchstone]")
        if ext_guidewire:
            lines.append("EXT_Guidewire[Guidewire]")
        if ext_imageright:
            lines.append("EXT_ImageRight[ImageRight]")
        if ext_genai:
            lines.append("EXT_GenAI[GenAI Platform]")

        # Edges (canonical backbone)
        if ui and api:
            lines.append("UI_UserInterface --> APP_ApiGateway")
        if api and auth:
            lines.append("APP_ApiGateway --> SEC_Auth")
        if (auth or api) and core:
            from_node = "SEC_Auth" if auth else "APP_ApiGateway"
            lines.append(f"{from_node} --> BIZ_CoreServices")
        if core and repo:
            lines.append("BIZ_CoreServices --> DATA_Repositories")
        if repo and db:
            lines.append("DATA_Repositories --> DATA_Database")
        if lb and ui:
            lines.append("INFRA_LoadBalancer --> UI_UserInterface")
        if cache and (api or core):
            lines.append(f"INFRA_Cache --> {'APP_ApiGateway' if api else 'BIZ_CoreServices'}")

        # External edges to core
        for ext in [
            (ext_ping, "EXT_PING"), (ext_touchstone, "EXT_Touchstone"),
            (ext_guidewire, "EXT_Guidewire"), (ext_imageright, "EXT_ImageRight"),
            (ext_genai, "EXT_GenAI")
        ]:
            if ext[0] and core:
                lines.append(f"{ext[1]} --> BIZ_CoreServices")

        # Styling
        lines += [
            "classDef UI fill:#E3F2FD,stroke:#1E88E5,stroke-width:2,color:#0D47A1;",
            "classDef APP fill:#E8F5E9,stroke:#43A047,stroke-width:2,color:#1B5E20;",
            "classDef BIZ fill:#FFF8E1,stroke:#FB8C00,stroke-width:2,color:#E65100;",
            "classDef DATA fill:#F3E5F5,stroke:#8E24AA,stroke-width:2,color:#4A148C;",
            "classDef SEC fill:#FFEBEE,stroke:#E53935,stroke-width:2,color:#B71C1C;",
            "classDef INFRA fill:#ECEFF1,stroke:#607D8B,stroke-width:2,color:#37474F;",
            "classDef EXT fill:#FFECB3,stroke:#F9A825,stroke-width:2,color:#E65100;",
        ]

        class_map = []
        if ui: class_map.append("class UI_UserInterface UI;")
        if api: class_map.append("class APP_ApiGateway APP;")
        if auth: class_map.append("class SEC_Auth SEC;")
        if core: class_map.append("class BIZ_CoreServices BIZ;")
        if repo: class_map.append("class DATA_Repositories DATA;")
        if db: class_map.append("class DATA_Database DATA;")
        if lb: class_map.append("class INFRA_LoadBalancer INFRA;")
        if cache: class_map.append("class INFRA_Cache INFRA;")
        if ext_ping: class_map.append("class EXT_PING EXT;")
        if ext_touchstone: class_map.append("class EXT_Touchstone EXT;")
        if ext_guidewire: class_map.append("class EXT_Guidewire EXT;")
        if ext_imageright: class_map.append("class EXT_ImageRight EXT;")
        if ext_genai: class_map.append("class EXT_GenAI EXT;")
        lines += class_map

    else:
        # LLD minimal from plan
        ctrl = has("controller", "resource", "endpoint")
        svc = has("service", "logic")
        rep = has("repository", "repo", "dao")
        dbtbl = has("table", "schema", "entity", "database")
        util_val = has("validate", "validation")
        util_err = has("error", "exception")

        if ctrl: lines.append("CTRL_Controller[Controller]")
        if svc: lines.append("SVC_Service[Service]")
        if rep: lines.append("REPO_Repository[Repository]")
        if dbtbl: lines.append("DB_Table[(Table)]")
        if util_val: lines.append("UTIL_Validation[Validation]")
        if util_err: lines.append("UTIL_Error[Error Handler]")

        if ctrl and svc: lines.append("CTRL_Controller --> SVC_Service")
        if svc and rep: lines.append("SVC_Service --> REPO_Repository")
        if rep and dbtbl: lines.append("REPO_Repository --> DB_Table")
        if util_val and svc: lines.append("SVC_Service --> UTIL_Validation")
        if util_err and svc: lines.append("SVC_Service --> UTIL_Error")

        lines += [
            "classDef CTRL fill:#E3F2FD,stroke:#1E88E5,stroke-width:2;",
            "classDef SVC fill:#E8F5E9,stroke:#43A047,stroke-width:2;",
            "classDef REPO fill:#FFF8E1,stroke:#FB8C00,stroke-width:2;",
            "classDef DB fill:#F3E5F5,stroke:#8E24AA,stroke-width:2;",
            "classDef UTIL fill:#ECEFF1,stroke:#607D8B,stroke-width:2;",
        ]
        class_map = []
        if ctrl: class_map.append("class CTRL_Controller CTRL;")
        if svc: class_map.append("class SVC_Service SVC;")
        if rep: class_map.append("class REPO_Repository REPO;")
        if dbtbl: class_map.append("class DB_Table DB;")
        if util_val: class_map.append("class UTIL_Validation UTIL;")
        if util_err: class_map.append("class UTIL_Error UTIL;")
        lines += class_map

    return "\n".join(lines) + "\n"

def agent_backlog_creator(plan, original_text, trd):
    """Enhanced backlog creator with better structure and fallback mechanism"""
    print("AGENT [Backlog Creator]: Creating comprehensive project backlog...")
    prompt = f"""You are a Senior Product Manager and Agile Coach with 15+ years of experience in software development.

Create a comprehensive project backlog based on the following plan, requirements, and technical requirements document. Structure it as a JSON object with the following format:

{{
  "backlog": [
    {{
      "id": "epic-1",
      "type": "Epic",
      "title": "Epic Title",
      "description": "Epic description",
      "priority": "High/Medium/Low",
      "effort": "Story points estimate",
      "trd_sections": ["Section 1", "Section 2"],
      "requirements_covered": ["Requirement 1", "Requirement 2"],
      "children": [
        {{
          "id": "feature-1",
          "type": "Feature",
          "title": "Feature Title",
          "description": "Feature description",
          "priority": "High/Medium/Low",
          "effort": "Story points estimate",
          "trd_sections": ["Section 1"],
          "requirements_covered": ["Requirement 1"],
          "children": [
            {{
              "id": "story-1",
              "type": "User Story",
              "title": "User Story Title",
              "description": "As a [user], I want [feature] so that [benefit]",
              "priority": "High/Medium/Low",
              "effort": "Story points estimate",
              "acceptance_criteria": ["Criterion 1", "Criterion 2"],
              "trd_sections": ["Section 1"],
              "requirements_covered": ["Requirement 1"]
            }}
          ]
        }}
      ]
    }}
  ]
}}

**Backlog Structure:**
1. **Epics**: High-level business capabilities (3-5 epics recommended)
2. **Features**: Major functionality within epics (2-4 features per epic)
3. **User Stories**: Detailed requirements with acceptance criteria (3-8 stories per feature)

**Linking Requirements:**
- For each backlog item, identify which TRD sections it addresses
- For each backlog item, identify which specific requirements it covers
- Use the TRD section headers (e.g., "FUNCTIONAL REQUIREMENTS", "SECURITY CONSIDERATIONS")
- Reference specific requirements from the original text

**Considerations:**
- Prioritize by business value and technical dependencies
- Include technical debt items
- Consider security, performance, and scalability requirements
- Provide realistic effort estimates (1-13 story points)
- Include acceptance criteria for all user stories
- Cover all major system components from the TRD
- Include testing, deployment, and maintenance stories
- Ensure each item links to specific TRD sections and requirements

**Requirements:**
- Create at least 3-5 epics covering different aspects of the system
- Each epic should have 2-4 features
- Each feature should have 3-8 user stories
- Total backlog should have 20-50 user stories
- Include stories for non-functional requirements (security, performance, etc.)
- Every item must have trd_sections and requirements_covered arrays

IMPORTANT: Respond ONLY with valid JSON, no additional text or explanations.

--- HIGH-LEVEL PLAN ---
{plan}

--- TECHNICAL REQUIREMENTS DOCUMENT ---
{trd if trd else "Technical Requirements Document not available"}

--- ORIGINAL REQUIREMENTS TEXT ---
{original_text}
"""
    result = call_generative_agent([{"text": prompt}], is_json=True, stage_name="backlog_creation")
    if result is None:
        print("AGENT [Backlog Creator]: API failed, creating fallback backlog...")
        # Create a fallback backlog structure with linking information
        fallback_backlog = {
            "backlog": [
                {
                    "id": "epic-1",
                    "type": "Epic",
                    "title": "System Implementation",
                    "description": "Core system implementation based on requirements analysis",
                    "priority": "High",
                    "effort": "40",
                    "trd_sections": ["SYSTEM OVERVIEW", "FUNCTIONAL REQUIREMENTS"],
                    "requirements_covered": ["System architecture setup", "Core functionality implementation"],
                    "children": [
                        {
                            "id": "feature-1",
                            "type": "Feature",
                            "title": "Basic System Setup",
                            "description": "Initial system setup and configuration",
                            "priority": "High",
                            "effort": "13",
                            "trd_sections": ["SYSTEM OVERVIEW"],
                            "requirements_covered": ["System architecture setup"],
                            "children": [
                                {
                                    "id": "story-1",
                                    "type": "User Story",
                                    "title": "System Architecture Setup",
                                    "description": "As a developer, I want to set up the system architecture so that the foundation is ready for development",
                                    "priority": "High",
                                    "effort": "5",
                                    "acceptance_criteria": ["Architecture diagram is created", "Technology stack is defined"],
                                    "trd_sections": ["SYSTEM OVERVIEW"],
                                    "requirements_covered": ["System architecture setup"]
                                }
                            ]
                        }
                    ]
                }
            ]
        }
        return json.dumps(fallback_backlog), None
    
    # Debug: Log the AI result
    print(f"AGENT [Backlog Creator]: AI returned result type: {type(result)}")
    print(f"AGENT [Backlog Creator]: AI result keys: {list(result.keys()) if isinstance(result, dict) else 'Not a dict'}")
    if isinstance(result, dict) and 'backlog' in result:
        print(f"AGENT [Backlog Creator]: AI backlog has {len(result['backlog'])} items")
        if result['backlog']:
            first_item = result['backlog'][0]
            print(f"AGENT [Backlog Creator]: First item title: {first_item.get('title', 'No title')}")
            print(f"AGENT [Backlog Creator]: First item type: {first_item.get('type', 'No type')}")
    
    return result, None

# --- Helper functions (add_ids_to_backlog, Azure DevOps, etc.) ---
def add_ids_to_backlog(items):
    if not isinstance(items, list): return []
    for i, epic in enumerate(items):
        epic['id'] = f"E-{i+1}"
        for j, feature in enumerate(epic.get('children', []), 1):
            feature['id'] = f"F-{j}"
            for k, story in enumerate(feature.get('children', []), 1):
                story['id'] = f"US-{k}"
    return items

# --- Enhanced Azure DevOps Integration ---

def get_ado_headers():
    """Get Azure DevOps headers with enhanced error handling"""
    if not ADO_PAT:
        raise ValueError("Azure DevOps Personal Access Token not configured")
    return {
        'Authorization': f'Basic {base64.b64encode(f":{ADO_PAT}".encode()).decode()}',
        'Content-Type': 'application/json-patch+json'
    }

def create_ado_work_item(item_type, title, description="", parent_url=None, tags=None):
    """Enhanced Azure DevOps work item creation with better error handling"""
    try:
        headers = get_ado_headers()
        
        # Prepare work item data
        work_item_data = [
            {
                "op": "add",
                "path": "/fields/System.Title",
                "value": title
            },
            {
                "op": "add",
                "path": "/fields/System.Description",
                "value": description
            }
        ]
        
        # Add tags if provided
        if tags:
            work_item_data.append({
                "op": "add",
                "path": "/fields/System.Tags",
                "value": "; ".join(tags)
            })
        
        # Add parent relationship if provided
        if parent_url:
            work_item_data.append({
                "op": "add",
                "path": "/relations/-",
                "value": {
                    "rel": "System.LinkTypes.Hierarchy-Reverse",
                    "url": parent_url
                }
            })
        
        url = f"{ADO_ORGANIZATION_URL}{ADO_PROJECT_NAME}/_apis/wit/workitems/${item_type}?api-version=6.0"
        
        response = requests.post(url, headers=headers, json=work_item_data)
        response.raise_for_status()
        
        result = response.json()
        print(f"Successfully created {item_type} work item: {result.get('id')}")
        return result
        
    except requests.exceptions.RequestException as e:
        print(f"Error creating Azure DevOps work item: {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return None
    except Exception as e:
        print(f"Unexpected error in Azure DevOps integration: {e}")
        return None

def process_backlog_for_ado(backlog):
    """Enhanced backlog processing for Azure DevOps with better structure"""
    if not ADO_PAT or not ADO_ORGANIZATION_URL or not ADO_PROJECT_NAME:
        print("Azure DevOps not configured, skipping work item creation")
        return {"success": False, "message": "Azure DevOps not configured"}
    
    try:
        created_items = []
        
        for epic in backlog:
            # Create Epic
            epic_description = f"Epic: {epic.get('title', 'Untitled Epic')}\n\nGenerated by BA Agent"
            epic_result = create_ado_work_item(
                "Epic", 
                epic.get('title', 'Untitled Epic'),
                epic_description,
                tags=["BA-Agent-Generated"]
            )
            
            if epic_result:
                epic_url = epic_result.get('url')
                created_items.append({
                    "type": "Epic",
                    "id": epic_result.get('id'),
                    "title": epic.get('title'),
                    "url": epic_url
                })
                
                # Create Features under Epic
                for feature in epic.get('children', []):
                    feature_description = f"Feature: {feature.get('title', 'Untitled Feature')}\n\nParent Epic: {epic.get('title')}"
                    feature_result = create_ado_work_item(
                        "Feature",
                        feature.get('title', 'Untitled Feature'),
                        feature_description,
                        parent_url=epic_url,
                        tags=["BA-Agent-Generated"]
                    )
                    
                    if feature_result:
                        feature_url = feature_result.get('url')
                        created_items.append({
                            "type": "Feature",
                            "id": feature_result.get('id'),
                            "title": feature.get('title'),
                            "url": feature_url
                        })
                        
                        # Create User Stories under Feature
                        for story in feature.get('children', []):
                            story_description = f"User Story: {story.get('title', 'Untitled Story')}\n\nParent Feature: {feature.get('title')}"
                            story_result = create_ado_work_item(
                                "User Story",
                                story.get('title', 'Untitled Story'),
                                story_description,
                                parent_url=feature_url,
                                tags=["BA-Agent-Generated"]
                            )
                            
                            if story_result:
                                created_items.append({
                                    "type": "User Story",
                                    "id": story_result.get('id'),
                                    "title": story.get('title'),
                                    "url": story_result.get('url')
                                })
        
        return {
            "success": True,
            "message": f"Successfully created {len(created_items)} work items in Azure DevOps",
            "items": created_items
        }
        
    except Exception as e:
        print(f"Error processing backlog for Azure DevOps: {e}")
        return {"success": False, "message": f"Error: {str(e)}"}

# --- Enhanced Email Notification System ---

def send_email_notification(subject, content, recipient_email=None, attachment_data=None):
    """Enhanced email notification system with better error handling and attachment support"""
    try:
        if not ACS_CONNECTION_STRING or not ACS_SENDER_ADDRESS:
            print("Email service not configured, skipping notification")
            return {"success": False, "message": "Email service not configured"}
        
        # Use provided recipient or default
        recipient = recipient_email or APPROVAL_RECIPIENT_EMAIL
        if not recipient:
            print("No recipient email configured")
            return {"success": False, "message": "No recipient email configured"}
        
        # Create email client
        email_client = EmailClient.from_connection_string(ACS_CONNECTION_STRING)
        
        # Prepare email content
        email_content = {
            "subject": subject,
            "plainText": content,
            "html": content.replace('\n', '<br>')
        }
        
        # Create message
        message = {
            "senderAddress": ACS_SENDER_ADDRESS,
            "recipients": {
                "to": [{"address": recipient}]
            },
            "content": email_content
        }
        
        # Add attachment if provided - using proper Azure Communication Services format
        if attachment_data:
            # Ensure proper attachment format for Azure Communication Services
            # According to Azure docs, attachments should have this structure
            attachment = {
                "name": attachment_data["name"],
                "contentType": attachment_data["contentType"],
                "content": attachment_data["content"]
            }
            message["attachments"] = [attachment]
            print(f"📎 Attachment added: {attachment_data['name']} ({len(attachment_data['content'])} chars)")
            
            # Also try alternative attachment format if the first doesn't work
            # Some Azure services expect different structure
            print(f"📎 Attachment structure: {attachment}")
        
        # Send email
        poller = email_client.begin_send(message)
        result = poller.result()
        
        print(f"✅ Email notification sent successfully to {recipient}")
        if attachment_data:
            print(f"📎 Email included attachment: {attachment_data['name']}")
        return {"success": True, "message": "Email sent successfully"}
        
    except Exception as e:
        print(f"❌ Error sending email notification: {e}")
        return {"success": False, "message": f"Error: {str(e)}"}

def send_analysis_completion_notification(analysis_id, filename, results_summary):
    """Send notification when analysis is completed"""
    subject = f"BA Agent Analysis Completed - {filename}"
    content = f"""
Analysis completed successfully!

File: {filename}
Analysis ID: {analysis_id}

Summary:
{results_summary}

You can view the full results in the BA Agent application.
"""
    
    return send_email_notification(subject, content)

def send_approval_notification(analysis_id, filename, approval_url, trd_content=None, trd_document_number=None):
    """Send notification for approval requests (simplified version without attachments)."""
    subject = f"BA Agent Analysis Ready for Approval - {filename}"
    content = f"""
A new analysis is ready for your approval.

File: {filename}
Analysis ID: {analysis_id}

Please review and approve the analysis at:
{approval_url}

This analysis includes:
- Technical Requirements Document
- High and Level Design Diagrams
- Project Backlog
- Risk Assessment
- Cost Estimation
- Sentiment Analysis
"""
    if trd_document_number:
        content += f"\nTRD Document Number: {trd_document_number}\n"
    
    # For now, skip attachments to ensure email delivery works
    # TODO: Re-enable attachments once table formatting is stable
    print(f"📧 Sending approval notification for analysis: {analysis_id}")
    print(f"📧 Subject: {subject}")
    print(f"📧 Content length: {len(content)} characters")
    
    return send_email_notification(subject, content)

def extract_mermaid_code(text):
    if not text:
        return ""
    match = re.search(r"```mermaid\n([\s\S]*?)```", text)
    if match:
        return match.group(1).strip()
    # Remove any standalone ```
    text = re.sub(r"```", "", text)
    return text.strip()

def render_mermaid_to_png(mermaid_code: str) -> BytesIO:
    """Render Mermaid code to PNG with error handling and code cleaning"""
    
    def clean_mermaid_code(code):
        """Clean problematic sequences conservatively to preserve Mermaid syntax"""
        if not code:
            return ""

        cleaned = code

        # Step 1: Replace HTML line breaks that often appear in labels
        cleaned = cleaned.replace('<br>', ' ')
        cleaned = cleaned.replace('<br/>', ' ')
        cleaned = cleaned.replace('<br />', ' ')

        # Step 2: Simplify noisy labels while keeping bracket structure intact
        import re
        cleaned = re.sub(r'([A-Z])\[([^\]]*?)(?:<br>|\([^)]*\))[^\]]*?\]',
                         lambda m: m.group(1) + '[' + m.group(2).strip() + ']', cleaned)

        # Step 3: Normalize patterns like Q[/policies (GET)] -> Q[policies]
        cleaned = re.sub(r'([A-Z])\[([^\]]*?)\/\([^)]*\)([^\]]*?)\]',
                         lambda m: m.group(1) + '[' + (m.group(2) + m.group(3)).strip() + ']', cleaned)

        # Do NOT strip critical characters like '<', '>', '/', '\\' as they are used in arrows and syntax
        # Preserve whitespace and newlines for Mermaid parser
        cleaned = cleaned.strip()

        return cleaned
    
    # Try original code first
    url = "https://kroki.io/mermaid/png"
    headers = {"Content-Type": "text/plain"}
    
    try:
        response = requests.post(url, data=mermaid_code.encode("utf-8"), headers=headers, timeout=30)
        if response.status_code == 200:
            return BytesIO(response.content)
    except Exception as e:
        print(f"Failed to render original Mermaid code: {e}")
    
    # Try with cleaned code as fallback
    try:
        cleaned_code = clean_mermaid_code(mermaid_code)
        if cleaned_code and cleaned_code != mermaid_code:
            print(f"Attempting to render cleaned Mermaid code...")
            response = requests.post(url, data=cleaned_code.encode("utf-8"), headers=headers, timeout=30)
            if response.status_code == 200:
                return BytesIO(response.content)
    except Exception as e:
        print(f"Failed to render cleaned Mermaid code: {e}")
    
    # If all else fails, create a placeholder image
    raise ValueError(f"Kroki diagram generation failed for both original and cleaned code")

# --- API Endpoints ---

@app.route("/api/test_image_extraction", methods=['POST'])
def test_image_extraction():
    """Test endpoint to check image extraction from documents"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part in the request"}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        print(f"🧪 Testing image extraction from: {file.filename}")
        
        # Test the extraction function
        text_content, images, error = agent_extract_content(io.BytesIO(file.read()), file.filename)
        
        if error:
            return jsonify({
                "success": False,
                "error": error,
                "filename": file.filename
            })
        
        return jsonify({
            "success": True,
            "filename": file.filename,
            "text_length": len(text_content),
            "images_count": len(images),
            "images": [
                {
                    "mime_type": img.get('mime_type', 'unknown'),
                    "data_length": len(img.get('data', ''))
                } for img in images
            ],
            "text_preview": text_content[:500] + "..." if len(text_content) > 500 else text_content
        })
        
    except Exception as e:
        print(f"❌ Exception during image extraction test: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        })

@app.route("/api/test_gemini", methods=['GET'])
def test_gemini_api():
    """Test endpoint to check if Gemini API is working"""
    try:
        print("🧪 Testing Gemini API connection...")
        
        headers = {
            "Content-Type": "application/json",
            "x-goog-api-key": GEMINI_API_KEY
        }
        
        payload = {
            "contents": [{
                "role": "user",
                "parts": [{"text": "Hello, please respond with 'API is working' if you can see this message."}]
            }],
            "generationConfig": {
                "temperature": 0.1,
                "maxOutputTokens": 100
            }
        }
        
        print(f"🔑 Using API key: {GEMINI_API_KEY[:10]}...")
        print(f"🌐 API URL: {GEMINI_API_URL}")
        
        response = requests.post(GEMINI_API_URL, headers=headers, json=payload, timeout=30)
        
        print(f"📡 Response status: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            content = result.get("candidates", [{}])[0].get("content", {})
            parts = content.get("parts", [])
            response_text = parts[0].get("text", "") if parts else ""
            
            return jsonify({
                "success": True,
                "message": "API is working",
                "response": response_text,
                "status_code": response.status_code
            })
        else:
            print(f"❌ API Error: {response.status_code} - {response.text}")
            return jsonify({
                "success": False,
                "error": f"API Error {response.status_code}",
                "response_text": response.text,
                "status_code": response.status_code
            })
            
    except Exception as e:
        print(f"❌ Exception during API test: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e),
            "status_code": 500
        })

@app.route("/api/test_diagrams", methods=['POST'])
def test_diagram_generation():
    """Test endpoint to check HLD and LLD diagram generation"""
    try:
        data = request.get_json()
        diagram_type = data.get('type', 'HLD')  # HLD or LLD
        test_plan = data.get('plan', 'Create a simple web application with user authentication and database storage.')
        
        print(f"🧪 Testing {diagram_type} diagram generation...")
        
        # Test the diagrammer function
        result, error = agent_diagrammer(test_plan, diagram_type)
        
        if error:
            return jsonify({
                "success": False,
                "diagram_type": diagram_type,
                "error": error,
                "result": result
            })
        else:
            return jsonify({
                "success": True,
                "diagram_type": diagram_type,
                "result": result,
                "error": None
            })
            
    except Exception as e:
        print(f"❌ Exception during diagram test: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e),
            "diagram_type": "unknown"
        })

@app.route("/api/debug_diagrams", methods=['GET'])
def debug_diagram_generation():
    """Debug endpoint to check diagram generation status"""
    try:
        # Test with a simple plan
        test_plan = "Create a web application with user authentication, database storage, and API endpoints."
        
        print("🔍 Debugging diagram generation...")
        
        # Test HLD
        hld_result, hld_error = agent_diagrammer(test_plan, "HLD")
        
        # Test LLD
        lld_result, lld_error = agent_diagrammer(test_plan, "LLD")
        
        return jsonify({
            "hld": {
                "success": not hld_error,
                "result": hld_result,
                "error": hld_error
            },
            "lld": {
                "success": not lld_error,
                "result": lld_result,
                "error": lld_error
            },
            "test_plan": test_plan
        })
        
    except Exception as e:
        print(f"❌ Exception during diagram debug: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "error": str(e)
        })

@app.route("/")
def index():
    return "<h1>AI Business Analyst Backend is running!</h1>"

@app.route("/api/generate", methods=['POST'])
def orchestrator():
    print("\n--- ORCHESTRATOR: Enhanced generation task started ---")
    if 'file' not in request.files:
        print("No file part in the request")
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['file']
    
    # Read file content once and store it
    file_content_bytes = file.read()
    file_size = len(file_content_bytes)
    print(f"DEBUG: File size: {file_size} bytes")
    
    # Extract content from file
    text_content, images, error = agent_extract_content(io.BytesIO(file_content_bytes), file.filename)
    if error:
        print(f"Extraction error: {error}")
        return jsonify({"error": error}), 500

    print(f"DEBUG: Content extracted successfully, length: {len(text_content)} characters")

    # Save document to database first
    try:
        # Create uploads directory if it doesn't exist
        uploads_dir = "uploads"
        if not os.path.exists(uploads_dir):
            os.makedirs(uploads_dir)
            print(f"📁 Created uploads directory: {uploads_dir}")
        
        # Create document record
        doc_id = str(uuid.uuid4())
        file_path = f"{uploads_dir}/{doc_id}_{file.filename}"
        
        # Save file to disk
        with open(file_path, 'wb') as f:
            f.write(file_content_bytes)
        print(f"💾 [File] Saved file to disk: {file_path}")
        
        # Classify LOB for the document
        lob_info = classify_line_of_business(text_content)

        document_data = {
            "id": doc_id,
            "name": file.filename,
            "uploadDate": datetime.now().isoformat(),
            "fileType": file.filename.split('.')[-1].lower(),
            "size": file_size,
            "status": "uploaded",
            "lob": lob_info
        }
        
        print(f"DEBUG: Document data created with ID: {doc_id}")
        
        # Save to database
        db = next(get_db())
        try:
            print(f"💾 [DB] Saving document to database: {file.filename}")
            saved_doc = save_document_to_db(db, user_email="guest", file_name=file.filename, file_type=document_data["fileType"], file_path=file_path, file_content=text_content, meta=document_data, status="uploaded")
            
            if saved_doc is None:
                print("⚠️ Warning: Failed to save document to database, but continuing with analysis")
            else:
                print(f"✅ [DB] Document stored with ID: {doc_id}")

                # Add to vector database
                print(f"📡 [VectorDB] Adding to vector DB...")
                add_to_vector_db(
                    content=text_content,
                    meta={
                        "id": doc_id,
                        "name": file.filename,
                        "type": "document",
                        "upload_date": document_data["uploadDate"],
                        "lob": lob_info
                    },
                    collection_name="documents"
                )
                print(f"✅ [VectorDB] Added to vector DB")

        except Exception as db_error:
            print(f"❌ [DB] Failed to save document: {db_error}")
            # Continue with analysis even if document save fails
        finally:
            try:
                db.close()
            except Exception:
                pass
                
    except Exception as e:
        print(f"❌ Error saving document: {e}")
        # Continue with analysis even if document save fails

    # 1. Enhanced Planning Agent
    plan, error = agent_planner(text_content, images)
    if error:
        print(f"Planning Agent failed: {error}")
        return jsonify({"error": f"Planning Agent failed: {error}"}), 500

    # 2. Specialist Agents (Enhanced with better error handling)
    print("AGENT [Orchestrator]: Starting specialist agents...")
    
    # Pass LOB context into TRD to enforce P&C US/EU focus when available
    try:
        # Try to infer LOB from the latest uploaded document in DB for context
        inferred_lob = None
        try:
            db = next(get_db())
            docs = get_all_documents_from_db(db)
            if docs:
                # Most recent uploaded document carries 'meta' with lob if via upload flow
                inferred_lob = (docs[0].get('meta') or {}).get('lob')
        finally:
            try:
                db.close()
            except Exception:
                pass
    except Exception:
        inferred_lob = None

    # Initialize enhanced document generator with LLM engine
    class LLMWrapper:
        def __init__(self):
            pass
        
        def generate_response(self, prompt, is_json=False):
            return call_generative_agent([{"text": prompt}], is_json=is_json, stage_name="enhanced_generation")
    
    enhanced_doc_generator.llm_engine = LLMWrapper()
    
    # Generate TRD using existing agent (keep for now)
    trd, err_trd = agent_trd_writer(plan, text_content, lob_info=inferred_lob)
    print(f"AGENT [Orchestrator]: TRD generation {'✅ Success' if not err_trd else f'❌ Failed: {err_trd}'}")
    
    # Generate HLD using enhanced generator
    hld, err_hld = enhanced_doc_generator.generate_high_quality_hld(plan, text_content)
    print(f"AGENT [Orchestrator]: Enhanced HLD generation {'✅ Success' if not err_hld else f'❌ Failed: {err_hld}'}")
    
    # Generate LLD using enhanced generator
    lld, err_lld = enhanced_doc_generator.generate_high_quality_lld(plan, text_content, hld if not err_hld else "")
    print(f"AGENT [Orchestrator]: Enhanced LLD generation {'✅ Success' if not err_lld else f'❌ Failed: {err_lld}'}")
    
    # Generate backlog using enhanced generator
    backlog_json, err_backlog = enhanced_doc_generator.generate_high_quality_backlog(plan, text_content, trd if not err_trd else "")
    print(f"AGENT [Orchestrator]: Enhanced backlog generation {'✅ Success' if not err_backlog else f'❌ Failed: {err_backlog}'}")
    
    # Debug: Log what we received from the backlog creator
    print(f"AGENT [Orchestrator]: Backlog creator returned:")
    print(f"  Type: {type(backlog_json)}")
    print(f"  Error: {err_backlog}")
    if isinstance(backlog_json, dict):
        print(f"  Keys: {list(backlog_json.keys())}")
        if 'backlog' in backlog_json:
            print(f"  Backlog length: {len(backlog_json['backlog'])}")
    elif isinstance(backlog_json, str):
        print(f"  String length: {len(backlog_json)}")
        print(f"  First 200 chars: {backlog_json[:200]}...")

    print(f"AGENT [Orchestrator]: Agent results:")
    print(f"  TRD: {'✅ Success' if not err_trd else f'❌ Failed: {err_trd}'}")
    print(f"  HLD: {'✅ Success' if not err_hld else f'❌ Failed: {err_hld}'}")
    print(f"  LLD: {'✅ Success' if not err_lld else f'❌ Failed: {err_lld}'}")
    print(f"  Backlog: {'✅ Success' if not err_backlog else f'❌ Failed: {err_backlog}'}")

    # Count successful and failed agents
    successful_agents = sum([1 for err in [err_trd, err_hld, err_lld, err_backlog] if not err])
    total_agents = 4
    
    print(f"AGENT [Orchestrator]: {successful_agents}/{total_agents} agents completed successfully")
    
    # Continue with partial results if at least some agents succeeded
    if successful_agents == 0:
        print("AGENT [Orchestrator]: All specialist agents failed")
        return jsonify({"error": "All specialist agents failed. Please try again."}), 500
    elif successful_agents < total_agents:
        print(f"AGENT [Orchestrator]: Partial success - {successful_agents}/{total_agents} agents completed")
    
    # 3. Final Assembly with Enhanced Results
    try:
        # Handle backlog processing (enhanced generator returns dict directly)
        actual_backlog_list = []
        if not err_backlog and backlog_json:
            try:
                # Enhanced generator returns dictionary directly
                if isinstance(backlog_json, dict):
                    actual_backlog_list = backlog_json.get('backlog', [])
                    print("AGENT [Orchestrator]: Enhanced backlog generator returned dictionary, extracted backlog list")
                else:
                    # Fallback: try to parse as JSON string
                    backlog_data = json.loads(backlog_json)
                    actual_backlog_list = backlog_data.get('backlog', [])
                    print("AGENT [Orchestrator]: Parsed backlog from JSON string (fallback)")
                
                if not isinstance(actual_backlog_list, list):
                    print("Warning: Backlog is not a list, correcting.")
                    actual_backlog_list = []
                    
                print(f"AGENT [Orchestrator]: Backlog contains {len(actual_backlog_list)} epics")
                
            except Exception as e:
                print(f"Backlog processing error: {e}")
                actual_backlog_list = []
        else:
            print("Backlog creation failed, using empty backlog")
        
        # Process enhanced document results
        hld_clean = ""
        lld_clean = ""
        
        if not err_hld and hld:
            # Enhanced generator returns Mermaid diagrams, extract the code
            hld_clean = extract_mermaid_code(hld)
            print(f"AGENT [Orchestrator]: Enhanced HLD result: {'✅ Available' if hld_clean else '❌ Empty'}")
        else:
            print(f"AGENT [Orchestrator]: Enhanced HLD failed or empty: {err_hld}")
            
        if not err_lld and lld:
            # Enhanced generator returns Mermaid diagrams, extract the code
            lld_clean = extract_mermaid_code(lld)
            print(f"AGENT [Orchestrator]: Enhanced LLD result: {'✅ Available' if lld_clean else '❌ Empty'}")
        else:
            print(f"AGENT [Orchestrator]: Enhanced LLD failed or empty: {err_lld}")
        
        print(f"AGENT [Orchestrator]: Final assembly:")
        print(f"  TRD: {'✅ Available' if trd else '❌ Not available'}")
        print(f"  HLD: {'✅ Available' if hld_clean else '❌ Not available'}")
        print(f"  LLD: {'✅ Available' if lld_clean else '❌ Not available'}")
        print(f"  Backlog: {'✅ Available' if actual_backlog_list else '❌ Not available'}")

        # Generate a TRD document number for this analysis
        trd_document_number = f"TRD-{datetime.now().strftime('%Y%m%d')}-{str(uuid.uuid4())[:8].upper()}"

        final_response = {
            "trd": trd or "Technical Requirements Document could not be generated due to API error.",
            "hld": hld_clean or "High-Level Design document could not be generated due to API error.",
            "lld": lld_clean or "Low-Level Design document could not be generated due to API error.",
            "images": images,
            "backlog": add_ids_to_backlog(actual_backlog_list),
            "enhanced_plan": plan or "",
            "partial_success": successful_agents < total_agents,
            "successful_agents": successful_agents,
            "total_agents": total_agents,
            "trd_document_number": trd_document_number,
            "agent_errors": {
                "trd": err_trd,
                "hld": err_hld,
                "lld": err_lld,
                "backlog": err_backlog
            },
            "generation_quality": {
                "enhanced_generation": True,
                "quality_validation": True,
                "multiple_attempts": True,
                "context_aware_fallback": True
            }
        }
        
        # Debug: Print the final backlog structure
        print(f"AGENT [Orchestrator]: Final backlog structure:")
        print(f"  Type: {type(final_response['backlog'])}")
        print(f"  Length: {len(final_response['backlog']) if isinstance(final_response['backlog'], list) else 'Not a list'}")
        if final_response['backlog']:
            print(f"  First item type: {type(final_response['backlog'][0])}")
            print(f"  First item keys: {list(final_response['backlog'][0].keys()) if isinstance(final_response['backlog'][0], dict) else 'Not a dict'}")
        
        # Debug: Print a sample of the backlog JSON
        try:
            backlog_json_sample = json.dumps(final_response['backlog'][:1], indent=2) if final_response['backlog'] else "[]"
            print(f"AGENT [Orchestrator]: Sample backlog JSON (first item):")
            print(backlog_json_sample[:500] + "..." if len(backlog_json_sample) > 500 else backlog_json_sample)
        except Exception as e:
            print(f"AGENT [Orchestrator]: Error serializing backlog sample: {e}")
    except Exception as e:
        import traceback
        print(f"Error during final assembly: {e}")
        traceback.print_exc()
        return jsonify({"error": f"Error during final assembly: {e}"}), 500

    print("--- ORCHESTRATOR: All enhanced agents completed. Task finished. ---\n")
    
    # Save analysis results first
    try:
        analysis_id = save_analysis_results(final_response, text_content, file.filename if file else "requirements.txt", document_id=doc_id if 'doc_id' in locals() else None)
        final_response["analysis_id"] = analysis_id
        print(f"AGENT [Orchestrator]: Analysis saved with ID: {analysis_id}")
    except Exception as e:
        print(f"AGENT [Orchestrator]: Warning - Failed to save analysis results: {e}")
        analysis_id = "unknown"
        final_response["analysis_id"] = analysis_id
    
    # Log total token consumption for this analysis
    try:
        total_tokens = sum(log['tokens_used'] for log in token_consumption_logs)
        log_token_consumption("analysis_complete", total_tokens, "gemini-1.5-pro", {
            "total_stages": len(token_consumption_logs),
            "analysis_id": analysis_id,
            "filename": file.filename if file else "requirements.txt"
        })
    except Exception as e:
        print(f"AGENT [Orchestrator]: Warning - Failed to log token consumption: {e}")
    
    # Send completion notification (optional - don't fail if this fails)
    try:
        results_summary = f"""
Analysis completed for {file.filename if file else 'requirements.txt'}

Generated:
- Technical Requirements Document
- High and Low Level Design Diagrams  
- Project Backlog with {len(actual_backlog_list)} epics

TRD Document Number: {final_response.get('trd_document_number','N/A')}
Analysis ID: {analysis_id}
"""
        
        email_result = send_analysis_completion_notification(
            analysis_id, 
            file.filename if file else "requirements.txt",
            results_summary
        )
        
        if email_result["success"]:
            print("Email notification sent successfully")
        else:
            print(f"Email notification failed: {email_result['message']}")
    except Exception as e:
        print(f"AGENT [Orchestrator]: Warning - Failed to send email notification: {e}")
    
    print("AGENT [Orchestrator]: Returning successful response to client")
    return jsonify(final_response)

@app.route("/api/render_mermaid", methods=["POST"])
def render_mermaid():
    data = request.get_json()
    mermaid_code = data.get("code")
    # Support Lucid handoff: if code is a Lucid token, return 204 so UI switches to embed
    if isinstance(mermaid_code, str) and mermaid_code.startswith("LUCID_EMBED::"):
        return ("", 204)
    if not mermaid_code:
        return jsonify({"error": "No Mermaid code provided"}), 400
    
    try:
        png_bytes = render_mermaid_to_png(mermaid_code)
        png_bytes.seek(0)
        return send_file(png_bytes, mimetype="image/png")
    except Exception as e:
        print(f"Error rendering Mermaid diagram: {e}")
        # Return a more detailed error response
        return jsonify({
            "error": f"Failed to render diagram: {str(e)}",
            "mermaid_code_preview": mermaid_code[:200] + "..." if len(mermaid_code) > 200 else mermaid_code
        }), 500

@app.route("/api/convert_to_docx", methods=['POST'])
def convert_to_docx():
    """Convert markdown content to DOCX format"""
    try:
        data = request.get_json()
        markdown_content = data.get('markdown', '')
        
        if not markdown_content:
            return jsonify({"error": "No markdown content provided"}), 400
        
        # Convert markdown to DOCX
        doc = markdown_to_docx(markdown_content)
        
        # Save to BytesIO
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return send_file(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='Technical_Requirements_Document.docx'
        )
        
    except Exception as e:
        print(f"Error converting to DOCX: {e}")
        return jsonify({"error": f"Failed to convert to DOCX: {str(e)}"}), 500

@app.route("/api/approve", methods=['POST'])
def send_for_approval():
    """Enhanced approval workflow with database storage"""
    try:
        data = request.get_json()
        analysis_id = data.get('analysis_id')
        results = data.get('results')
        
        if not analysis_id or not results:
            return jsonify({"error": "Missing analysis_id or results"}), 400
        
        # Generate unique approval ID
        approval_id = str(uuid.uuid4())
        
        # Create approval record with enhanced status tracking
        approval_record = {
            "id": approval_id,
            "analysis_id": analysis_id,
            "status": "pending",
            "created_date": datetime.now().isoformat(),  # Changed from created_at
            "updated_date": datetime.now().isoformat(),  # Changed from updated_at
            "approver_email": APPROVAL_RECIPIENT_EMAIL,
            "results_summary": {
                "has_trd": bool(results.get('trd')),
                "has_diagrams": bool(results.get('hld') or results.get('lld')),
                "has_backlog": bool(results.get('backlog')),
                "backlog_epics_count": len(results.get('backlog', []))
            }
        }
        
        # Store approval record in database
        try:
            db = next(get_db())
            save_approval_to_db(db, approval_record)
            print(f"✅ Approval saved to database: {approval_id}")
        except Exception as e:
            print(f"❌ Error saving approval to database: {e}")
            # Fallback to in-memory storage
            approval_statuses[approval_id] = approval_record
            print(f"⚠️ Using in-memory fallback for approval: {approval_id}")
        
        # Generate approval URL
        approval_url = f"{BACKEND_BASE_URL}/api/approval_response?approval_id={approval_id}"
        
        # Send approval notification email (optional - don't fail if this fails)
        try:
            trd_document_number = results.get('trd_document_number', 'N/A')
            trd_content = results.get('trd', None)  # Get TRD content for attachment
            
            # Debug: Log TRD content status
            if trd_content:
                print(f"📄 TRD content found: {len(trd_content)} characters")
                print(f"📄 TRD preview: {trd_content[:200]}...")
            else:
                print(f"⚠️ No TRD content found in results")
                print(f"📄 Available keys in results: {list(results.keys())}")
            
            email_result = send_approval_notification(analysis_id, "Analysis Results", approval_url, trd_content=trd_content, trd_document_number=trd_document_number)
            
            if email_result["success"]:
                print(f"✅ Approval notification sent successfully for approval ID: {approval_id}")
            else:
                print(f"⚠️ Approval notification failed: {email_result['message']} - continuing with approval workflow")
        except Exception as e:
            print(f"⚠️ Email notification error (continuing): {e}")
        
        return jsonify({
            "success": True,
            "approval_id": approval_id,
            "message": "Analysis sent for approval successfully",
            "approval_url": approval_url
        })
        
    except Exception as e:
        print(f"❌ Error in approval workflow: {e}")
        return jsonify({"error": f"Failed to send for approval: {str(e)}"}), 500

@app.route("/api/convert-to-docx", methods=["POST"])
def convert_markdown_to_docx():
    """Convert markdown content to DOCX format"""
    try:
        data = request.get_json()
        if not data or 'content' not in data:
            return jsonify({"error": "Content is required"}), 400
        
        content = data['content']
        filename = data.get('filename', 'document.docx')
        
        # Convert markdown to DOCX
        doc = markdown_to_docx(content)
        
        # Create response with DOCX file
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"❌ Error converting markdown to DOCX: {e}")
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

@app.route("/api/test-trd-attachment", methods=["POST"])
def test_trd_attachment():
    """Test endpoint to verify TRD attachment creation"""
    try:
        data = request.get_json()
        if not data or 'trd_content' not in data:
            return jsonify({"error": "TRD content is required"}), 400
        
        trd_content = data['trd_content']
        analysis_id = data.get('analysis_id', 'test')
        
        print(f"🧪 Testing TRD attachment creation for analysis: {analysis_id}")
        print(f"📄 TRD content length: {len(trd_content)} characters")
        
        # Test the attachment creation process
        try:
            # Convert TRD content to DOCX
            doc = markdown_to_docx(trd_content)
            
            # Save to BytesIO
            docx_bytes = BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            # Create attachment data
            docx_content = docx_bytes.getvalue()
            base64_content = base64.b64encode(docx_content).decode('utf-8')
            
            attachment_data = {
                "name": f"Test_TRD_{analysis_id}.docx",
                "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "content": base64_content
            }
            
            print(f"✅ Test TRD attachment created successfully")
            print(f"📎 File size: {len(docx_content)} bytes")
            print(f"📎 Base64 length: {len(base64_content)} characters")
            
            return jsonify({
                "success": True,
                "message": "TRD attachment test successful",
                "attachment_info": {
                    "name": attachment_data["name"],
                    "size_bytes": len(docx_content),
                    "base64_length": len(base64_content),
                    "content_type": attachment_data["contentType"]
                }
            })
            
        except Exception as e:
            print(f"❌ Error in TRD attachment creation: {e}")
            return jsonify({"error": f"TRD attachment creation failed: {str(e)}"}), 500
        
    except Exception as e:
        print(f"❌ Error in TRD attachment test: {e}")
        return jsonify({"error": f"Test failed: {str(e)}"}), 500

@app.route("/api/save-trd", methods=["POST"])
def save_trd_document():
    """Save TRD document with editing capabilities"""
    try:
        data = request.get_json()
        if not data or 'trd_content' not in data:
            return jsonify({"error": "TRD content is required"}), 400
        
        analysis_id = data.get('analysis_id')
        trd_content = data['trd_content']
        title = data.get('title', 'Technical Requirements Document')
        timestamp = data.get('timestamp')
        
        if not analysis_id:
            return jsonify({"error": "Analysis ID is required"}), 400
        
        # Get database connection
        db = next(get_db())
        
        try:
            # Get the existing analysis
            analysis = db.query(Analysis).filter(Analysis.id == analysis_id).first()
            if not analysis:
                return jsonify({"error": "Analysis not found"}), 404
            
            # Update the TRD content in the results
            if analysis.results:
                results = analysis.results
            else:
                results = {}
            
            results['trd'] = trd_content
            results['trd_last_updated'] = timestamp or datetime.now().isoformat()
            
            # Update the analysis in the database
            analysis.results = results
            db.commit()
            
            print(f"✅ TRD document saved successfully for analysis: {analysis_id}")
            
            return jsonify({
                "success": True,
                "message": "TRD document saved successfully",
                "analysis_id": analysis_id,
                "timestamp": timestamp or datetime.now().isoformat()
            })
            
        except Exception as db_error:
            db.rollback()
            print(f"❌ Database error saving TRD: {db_error}")
            return jsonify({"error": f"Database save failed: {str(db_error)}"}), 500
        finally:
            db.close()
        
    except Exception as e:
        print(f"❌ Error saving TRD document: {e}")
        return jsonify({"error": f"Save failed: {str(e)}"}), 500

@app.route("/api/approval_response")
def handle_approval_response():
    """Enhanced approval response handler with database storage"""
    try:
        approval_id = request.args.get('approval_id')
        action = request.args.get('action', 'view')  # view, approve, reject
        
        # Try to get approval from database first
        db = next(get_db())
        approval_record = get_approval_from_db(db, approval_id)
        
        # Fallback to in-memory storage if not in database
        if not approval_record and approval_id in approval_statuses:
            approval_record = approval_statuses[approval_id]
        
        if not approval_record:
            return jsonify({"error": "Invalid approval ID"}), 404
        
        if action == 'view':
            # Show approval interface
            return f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>BA Agent - Analysis Approval</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; }}
                    .container {{ max-width: 800px; margin: 0 auto; }}
                    .header {{ background: #f8f9fa; padding: 20px; border-radius: 8px; margin-bottom: 20px; }}
                    .summary {{ background: #e9ecef; padding: 15px; border-radius: 8px; margin-bottom: 20px; }}
                    .actions {{ display: flex; gap: 10px; margin-top: 20px; }}
                    .btn {{ padding: 10px 20px; border: none; border-radius: 4px; cursor: pointer; text-decoration: none; }}
                    .btn-approve {{ background: #28a745; color: white; }}
                    .btn-reject {{ background: #dc3545; color: white; }}
                    .btn-view {{ background: #007bff; color: white; }}
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="header">
                        <h1>BA Agent Analysis Approval</h1>
                        <p><strong>Approval ID:</strong> {approval_id}</p>
                        <p><strong>Analysis ID:</strong> {approval_record['analysis_id']}</p>
                        <p><strong>Created:</strong> {approval_record['created_date']}</p>
                    </div>
                    
                    <div class="summary">
                        <h2>Analysis Summary</h2>
                        <ul>
                            <li>Technical Requirements Document: {'✓' if approval_record['results_summary']['has_trd'] else '✗'}</li>
                            <li>Design Diagrams: {'✓' if approval_record['results_summary']['has_diagrams'] else '✗'}</li>
                            <li>Project Backlog: {'✓' if approval_record['results_summary']['has_backlog'] else '✗'}</li>
                            <li>Backlog Epics: {approval_record['results_summary']['backlog_epics_count']}</li>
                        </ul>
                    </div>
                    
                    <div class="actions">
                        <a href="?approval_id={approval_id}&action=approve" class="btn btn-approve">Approve Analysis</a>
                        <a href="?approval_id={approval_id}&action=reject" class="btn btn-reject">Reject Analysis</a>
                        <a href="/" class="btn btn-view">View in Application</a>
                    </div>
                </div>
            </body>
            </html>
            """
        
        elif action in ['approve', 'reject']:
            # Update approval status
            update_data = {
                'status': action,
                'updated_date': datetime.now().isoformat(),
                'approver_response': action
            }
            
            # Try to update in database first
            db_updated = update_approval_in_db_with_data(db, approval_id, update_data)
            
            # Fallback to in-memory update
            if not db_updated:
                approval_record['status'] = action
                approval_record['updated_date'] = datetime.now().isoformat()
                approval_record['approver_response'] = action
                approval_statuses[approval_id] = approval_record
            
            # If approved, automatically create ADO work items
            ado_result = None
            if action == 'approve':
                try:
                    # Get the analysis results to extract backlog
                    analysis_id = approval_record['analysis_id']
                    analysis = get_analysis_by_id(analysis_id)
                    
                    print(f"Processing approval for analysis {analysis_id}")
                    print(f"Analysis found: {analysis is not None}")
                    
                    if analysis and analysis.get('results', {}).get('backlog'):
                        backlog = analysis['results']['backlog']
                        print(f"Backlog found with {len(backlog)} epics")
                        
                        # Check Azure DevOps configuration first
                        if not ADO_PAT or not ADO_ORGANIZATION_URL or not ADO_PROJECT_NAME:
                            print("Azure DevOps not configured - skipping work item creation")
                            ado_result = {
                                "success": False, 
                                "message": "Azure DevOps not configured. Please set ADO_PERSONAL_ACCESS_TOKEN, ADO_ORGANIZATION_URL, and ADO_PROJECT_NAME environment variables."
                            }
                        else:
                            print("Azure DevOps configured - creating work items...")
                            ado_result = process_backlog_for_ado(backlog)
                            
                            if ado_result.get('success'):
                                items_created = len(ado_result.get('items', []))
                                print(f"Successfully created {items_created} ADO work items")
                                
                                # Log details of created items
                                for item in ado_result.get('items', []):
                                    print(f"  Created {item['type']}: {item['title']} (ID: {item['id']})")
                            else:
                                print(f"ADO work item creation failed: {ado_result.get('message')}")
                    else:
                        print(f"No backlog found for analysis {analysis_id}")
                        ado_result = {
                            "success": False, 
                            "message": "No backlog found in analysis results"
                        }
                except Exception as e:
                    print(f"Error creating ADO work items: {e}")
                    import traceback
                    traceback.print_exc()
                    ado_result = {"success": False, "message": f"Error: {str(e)}"}
            
            # Update ADO result in database
            if ado_result:
                update_approval_in_db_with_data(db, approval_id, {'ado_result': ado_result})
            
            # Send notification email
            subject = f"Analysis {'Approved' if action == 'approve' else 'Rejected'} - {approval_record['analysis_id']}"
            content = f"""
Analysis has been {action}ed.

Approval ID: {approval_id}
Analysis ID: {approval_record['analysis_id']}
Status: {action.upper()}
Updated: {approval_record['updated_date']}

You can view the full results in the BA Agent application.
"""
            
            # Add ADO information to email if approved
            if action == 'approve' and ado_result:
                if ado_result.get('success'):
                    content += f"\n\nADO Work Items Created: {len(ado_result.get('items', []))} items"
                    content += f"\nADO Status: Success"
                else:
                    content += f"\n\nADO Work Items: Failed to create"
                    content += f"\nADO Error: {ado_result.get('message')}"
            
            email_result = send_email_notification(subject, content)
            
            if email_result["success"]:
                print(f"✅ Approval {action} notification sent successfully")
            else:
                print(f"❌ Approval {action} notification failed: {email_result['message']}")
            
            return jsonify({
                "success": True,
                "status": action,
                "message": f"Analysis {action}ed successfully",
                "ado_result": ado_result
            })
        
        else:
            return jsonify({"error": "Invalid action"}), 400
            
    except Exception as e:
        print(f"❌ Error in approval response handler: {e}")
        return jsonify({"error": f"Failed to process approval response: {str(e)}"}), 500

@app.route("/api/approval_status/<approval_id>")
def get_approval_status(approval_id):
    """Get approval status with database storage"""
    try:
        # Try to get approval from database first
        db = next(get_db())
        approval_record = get_approval_from_db(db, approval_id)
        
        # Fallback to in-memory storage if not in database
        if not approval_record and approval_id in approval_statuses:
            approval_record = approval_statuses[approval_id]
        
        if not approval_record:
            return jsonify({"error": "Approval not found"}), 404
        
        # If approved, try to get Azure DevOps results
        ado_results = None
        if approval_record.get('status') == 'approve':
            try:
                # Get the analysis results to extract backlog
                analysis_id = approval_record['analysis_id']
                analysis = get_analysis_by_id(analysis_id)
                
                if analysis and analysis.get('results', {}).get('backlog'):
                    # Check if we already have ADO results in the approval record
                    if 'ado_result' in approval_record:
                        ado_results = approval_record['ado_result']
                    else:
                        # Try to create ADO work items if not already done
                        if ADO_PAT and ADO_ORGANIZATION_URL and ADO_PROJECT_NAME:
                            ado_results = process_backlog_for_ado(analysis['results']['backlog'])
                            # Store the results in the approval record
                            update_approval_in_db(db, approval_id, {'ado_result': ado_results})
                        else:
                            ado_results = {
                                "success": False,
                                "message": "Azure DevOps not configured"
                            }
            except Exception as e:
                ado_results = {
                    "success": False,
                    "message": f"Error processing Azure DevOps: {str(e)}"
                }
        
        return jsonify({
            **approval_record,
            "ado_result": ado_results
        })
        
    except Exception as e:
        print(f"❌ Error getting approval status: {e}")
        return jsonify({"error": f"Failed to get approval status: {str(e)}"}), 500

@app.route("/api/search", methods=['POST'])
def semantic_search():
    """Perform semantic search across documents and analyses"""
    try:
        data = request.get_json()
        query = data.get('query', '')
        collection = data.get('collection', 'documents')
        n_results = data.get('n_results', 5)
        restrict_pnc = data.get('restrict_pnc', True)
        
        if not query:
            return jsonify({"error": "Query is required"}), 400
        
        # Perform vector search
        results = search_vector_db(query, collection, n_results)

        # Optionally restrict results to US/Europe P&C Insurance based on stored LOB metadata
        if restrict_pnc and results and results.get('metas'):
            filtered_indices = []
            for idx, meta in enumerate(results.get('metas', [])):
                lob = (meta or {}).get('lob') or {}
                if (
                    lob.get('industry') == 'Insurance' and
                    lob.get('segment') == 'P&C' and
                    lob.get('region') in ['US', 'Europe', 'US/EU-unknown']
                ):
                    filtered_indices.append(idx)

            def apply_filter(seq):
                return [seq[i] for i in filtered_indices] if isinstance(seq, list) else seq

            results = {
                "documents": apply_filter(results.get('documents', [])),
                "metas": apply_filter(results.get('metas', [])),
                "distances": apply_filter(results.get('distances', [])),
                "ids": apply_filter(results.get('ids', []))
            }
        
        return jsonify({
            "query": query,
            "results": results,
            "total_results": len(results.get('documents', []))
        })
        
    except Exception as e:
        print(f"Error in semantic search: {e}")
        return jsonify({"error": f"Search failed: {str(e)}"}), 500

@app.route("/api/vector/collections", methods=['GET'])
def get_collections():
    """Get all Qdrant collections"""
    try:
        collections = qdrant_client.get_collections()
        return jsonify({
            "collections": [col.name for col in collections.collections]
        })
    except Exception as e:
        print(f"Error getting collections: {e}")
        return jsonify({"error": f"Failed to get collections: {str(e)}"}), 500

@app.route("/api/vector/collections/<collection_name>/info", methods=['GET'])
def get_collection_info(collection_name):
    """Get information about a specific collection"""
    try:
        collection_info = qdrant_client.get_collection(collection_name)
        return jsonify({
            "name": collection_name,
            "vectors_count": collection_info.vectors_count,
            "points_count": collection_info.points_count,
            "status": collection_info.status
        })
    except Exception as e:
        print(f"Error getting collection info: {e}")
        return jsonify({"error": f"Failed to get collection info: {str(e)}"}), 500

@app.route("/api/vector/collections/<collection_name>/points", methods=['GET'])
def get_collection_points(collection_name):
    """Get points from a collection"""
    try:
        limit = request.args.get('limit', 100, type=int)
        offset = request.args.get('offset', 0, type=int)
        
        points = qdrant_client.scroll(
            collection_name=collection_name,
            limit=limit,
            offset=offset
        )
        
        return jsonify({
            "collection": collection_name,
            "points": [
                {
                    "id": point.id,
                    "payload": point.payload,
                    "score": getattr(point, 'score', None)
                }
                for point in points[0]
            ],
            "total": len(points[0])
        })
    except Exception as e:
        print(f"Error getting collection points: {e}")
        return jsonify({"error": f"Failed to get collection points: {str(e)}"}), 500

@app.route("/api/vector/collections/<collection_name>/points/<point_id>", methods=['DELETE'])
def delete_point(collection_name, point_id):
    """Delete a point from a collection"""
    try:
        success = delete_from_vector_db(point_id, collection_name)
        if success:
            return jsonify({"message": f"Point {point_id} deleted successfully"})
        else:
            return jsonify({"error": "Failed to delete point"}), 500
    except Exception as e:
        print(f"Error deleting point: {e}")
        return jsonify({"error": f"Failed to delete point: {str(e)}"}), 500

# --- Enhanced Image Processing ---
def setup_enhanced_image_processing():
    """Setup enhanced image processing with multiple libraries"""
    print("🔧 Setting up enhanced image processing...")
    
    available_libraries = {
        'PIL/Pillow': False,
        'ImageMagick (Wand)': False,
        'OpenCV': False,
        'pdf2image': False
    }
    
    try:
        from PIL import Image, ImageDraw
        available_libraries['PIL/Pillow'] = True
        print("✅ PIL/Pillow available")
    except ImportError:
        print("❌ PIL/Pillow not available")
    
    try:
        from wand.image import Image as WandImage
        available_libraries['ImageMagick (Wand)'] = True
        print("✅ ImageMagick (Wand) available")
    except ImportError:
        print("❌ ImageMagick (Wand) not available")
    
    try:
        import cv2
        available_libraries['OpenCV'] = True
        print("✅ OpenCV available")
    except ImportError:
        print("❌ OpenCV not available")
    
    try:
        from pdf2image import convert_from_bytes
        available_libraries['pdf2image'] = True
        print("✅ pdf2image available")
    except ImportError:
        print("❌ pdf2image not available")
    
    print(f"📊 Available libraries: {sum(available_libraries.values())}/{len(available_libraries)}")
    return available_libraries

def process_images_with_multiple_methods(images):
    """Process images using multiple conversion methods"""
    processed_images = []
    
    for img in images:
        mime_type = img.get('mime_type', '')
        image_data = base64.b64decode(img.get('data', ''))
        
        # Method 1: Try PIL
        processed_img = convert_with_pil(image_data, mime_type)
        if processed_img:
            processed_images.append(processed_img)
            continue
        
        # Method 2: Try ImageMagick
        processed_img = convert_with_wand(image_data, mime_type)
        if processed_img:
            processed_images.append(processed_img)
            continue
        
        # Method 3: Try OpenCV
        processed_img = convert_with_opencv(image_data, mime_type)
        if processed_img:
            processed_images.append(processed_img)
            continue
        
        # Method 4: Create placeholder
        processed_img = create_placeholder(mime_type)
        if processed_img:
            processed_images.append(processed_img)
    
    return processed_images

def convert_with_pil(image_data, mime_type):
    """Convert image using PIL"""
    try:
        from PIL import Image
        import io
        
        image = Image.open(io.BytesIO(image_data))
        if image.mode in ('RGBA', 'LA', 'P'):
            image = image.convert('RGB')
        
        output_buffer = io.BytesIO()
        image.save(output_buffer, format='PNG')
        converted_data = output_buffer.getvalue()
        
        return {
            'mime_type': 'image/png',
            'data': base64.b64encode(converted_data).decode('utf-8')
        }
    except Exception as e:
        print(f"DEBUG: PIL conversion failed for {mime_type}: {e}")
        return None

def convert_with_wand(image_data, mime_type):
    """Convert image using ImageMagick"""
    try:
        from wand.image import Image as WandImage
        
        with WandImage(blob=image_data) as wand_img:
            wand_img.format = 'png'
            converted_data = wand_img.make_blob()
            
            return {
                'mime_type': 'image/png',
                'data': base64.b64encode(converted_data).decode('utf-8')
            }
    except ImportError:
        print("DEBUG: ImageMagick (wand) not available")
        return None
    except Exception as e:
        print(f"DEBUG: ImageMagick conversion failed for {mime_type}: {e}")
        return None

def convert_with_opencv(image_data, mime_type):
    """Convert image using OpenCV"""
    try:
        import cv2
        import numpy as np
        
        nparr = np.frombuffer(image_data, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if img is None:
            return None
        
        img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        success, buffer = cv2.imencode('.png', img_rgb)
        
        if success:
            converted_data = buffer.tobytes()
            return {
                'mime_type': 'image/png',
                'data': base64.b64encode(converted_data).decode('utf-8')
            }
    except ImportError:
        print("DEBUG: OpenCV not available")
        return None
    except Exception as e:
        print(f"DEBUG: OpenCV conversion failed for {mime_type}: {e}")
        return None

def create_placeholder(mime_type):
    """Create placeholder image"""
    try:
        from PIL import Image, ImageDraw
        
        placeholder = Image.new('RGB', (400, 300), color='lightgray')
        draw = ImageDraw.Draw(placeholder)
        
        draw.text((20, 20), f"Image: {mime_type}", fill='black')
        draw.text((20, 50), "Format not supported", fill='red')
        draw.text((20, 80), "Converted to placeholder", fill='blue')
        
        output_buffer = io.BytesIO()
        placeholder.save(output_buffer, format='PNG')
        converted_data = output_buffer.getvalue()
        
        return {
            'mime_type': 'image/png',
            'data': base64.b64encode(converted_data).decode('utf-8')
        }
    except Exception as e:
        print(f"DEBUG: Failed to create placeholder: {e}")
        return None

# Initialize enhanced image processing
ENHANCED_IMAGE_LIBRARIES = setup_enhanced_image_processing()

@app.route("/api/ado/status", methods=['GET'])
def check_ado_status():
    """Check Azure DevOps configuration and connection status"""
    try:
        # Check if Azure DevOps is configured
        if not ADO_PAT or not ADO_ORGANIZATION_URL or not ADO_PROJECT_NAME:
            return jsonify({
                "configured": False,
                "message": "Azure DevOps not configured",
                "missing": []
            })
        
        # Check what's missing
        missing = []
        if not ADO_PAT:
            missing.append("ADO_PERSONAL_ACCESS_TOKEN")
        if not ADO_ORGANIZATION_URL:
            missing.append("ADO_ORGANIZATION_URL")
        if not ADO_PROJECT_NAME:
            missing.append("ADO_PROJECT_NAME")
        
        if missing:
            return jsonify({
                "configured": False,
                "message": f"Missing configuration: {', '.join(missing)}",
                "missing": missing
            })
        
        # Test connection by trying to get project info
        try:
            headers = get_ado_headers()
            test_url = f"{ADO_ORGANIZATION_URL}{ADO_PROJECT_NAME}/_apis/project?api-version=6.0"
            response = requests.get(test_url, headers=headers)
            
            if response.status_code == 200:
                project_info = response.json()
                return jsonify({
                    "configured": True,
                    "connected": True,
                    "message": "Azure DevOps connection successful",
                    "project": {
                        "name": project_info.get('name'),
                        "id": project_info.get('id'),
                        "description": project_info.get('description')
                    }
                })
            else:
                return jsonify({
                    "configured": True,
                    "connected": False,
                    "message": f"Connection failed: {response.status_code} - {response.text}",
                    "error": response.text
                })
                
        except Exception as e:
            return jsonify({
                "configured": True,
                "connected": False,
                "message": f"Connection error: {str(e)}",
                "error": str(e)
            })
            
    except Exception as e:
        return jsonify({
            "configured": False,
            "connected": False,
            "message": f"Error checking status: {str(e)}",
            "error": str(e)
        })

@app.route("/api/ado/test", methods=['POST'])
def test_ado_creation():
    """Test Azure DevOps work item creation"""
    try:
        # Check if Azure DevOps is configured
        if not ADO_PAT or not ADO_ORGANIZATION_URL or not ADO_PROJECT_NAME:
            return jsonify({
                "success": False,
                "message": "Azure DevOps not configured"
            })
        
        # Create a test work item
        test_title = "BA Agent Test Work Item"
        test_description = "This is a test work item created by BA Agent to verify Azure DevOps integration."
        
        result = create_ado_work_item("Task", test_title, test_description, tags=["BA-Agent-Test"])
        
        if result:
            return jsonify({
                "success": True,
                "message": "Test work item created successfully",
                "work_item": {
                    "id": result.get('id'),
                    "url": result.get('url'),
                    "title": test_title
                }
            })
        else:
            return jsonify({
                "success": False,
                "message": "Failed to create test work item"
            })
            
    except Exception as e:
        return jsonify({
            "success": False,
            "message": f"Error testing Azure DevOps: {str(e)}"
        })

# ============================================================================
# MULTI-USER AUTHENTICATION ENDPOINTS
# ============================================================================

@app.route('/api/auth/register', methods=['POST'])
def api_register():
    """Register a new user"""
    return register_user()

@app.route('/api/auth/login', methods=['POST'])
def api_login():
    """Login user"""
    return login_user()

@app.route('/api/auth/logout', methods=['POST'])
@token_required
def api_logout():
    """Logout user"""
    return logout_user()

@app.route('/api/auth/profile', methods=['GET'])
@token_required
def api_get_profile():
    """Get user profile"""
    return get_user_profile()

@app.route('/api/auth/profile', methods=['PUT'])
@token_required
def api_update_profile():
    """Update user profile"""
    return update_user_profile()

@app.route('/api/auth/change-password', methods=['POST'])
@token_required
def api_change_password():
    """Change user password"""
    return change_password()

# Admin endpoints
@app.route('/api/admin/users', methods=['GET'])
@token_required
@role_required(['admin'])
def api_get_all_users():
    """Get all users (admin only)"""
    return get_all_users()

@app.route('/api/admin/users/status', methods=['PUT'])
@token_required
@role_required(['admin'])
def api_update_user_status():
    """Update user status (admin only)"""
    return update_user_status()

# ============================================================================
# PROJECT MANAGEMENT ENDPOINTS
# ============================================================================

@app.route('/api/projects', methods=['POST'])
@token_required
def create_project_api():
    """Create a new project"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'message': 'No data provided'}), 400
        
        name = data.get('name')
        description = data.get('description', '')
        
        if not name:
            return jsonify({'message': 'Project name is required'}), 400
        
        db = next(get_db())
        try:
            # Create project
            project = create_project(
                db=db,
                name=name,
                description=description,
                owner_id=request.current_user['id']
            )
            
            # Add owner to project members
            add_user_to_project(db, request.current_user['id'], project.id, ProjectUserRole.OWNER)
            
            return jsonify({
                'message': 'Project created successfully',
                'project': {
                    'id': project.id,
                    'name': project.name,
                    'description': project.description,
                    'owner_id': project.owner_id,
                    'status': project.status,
                    'created_at': project.created_at.isoformat()
                }
            }), 201
            
        finally:
            db.close()
            
    except Exception as e:
        return jsonify({'message': f'Failed to create project: {str(e)}'}), 500

@app.route('/api/projects', methods=['GET'])
@token_required
def get_projects_api():
    """Get all projects for current user"""
    try:
        db = next(get_db())
        try:
            projects = get_user_projects(db, request.current_user['id'])
            
            return jsonify({
                'projects': projects
            }), 200
            
        finally:
            db.close()
            
    except Exception as e:
        return jsonify({'message': f'Failed to get projects: {str(e)}'}), 500

@app.route('/api/projects/<project_id>', methods=['GET'])
@token_required
@project_access_required
def get_project_api(project_id):
    """Get specific project details"""
    try:
        db = next(get_db())
        try:
            project = db.query(Project).filter(Project.id == project_id).first()
            
            if not project:
                return jsonify({'message': 'Project not found'}), 404
            
            # Get project statistics
            doc_count = db.query(Document).join(
                project_documents
            ).filter(
                project_documents.c.project_id == project_id
            ).count()
            
            analysis_count = db.query(Analysis).filter(
                Analysis.project_id == project_id
            ).count()
            
            return jsonify({
                'project': {
                    'id': project.id,
                    'name': project.name,
                    'description': project.description,
                    'owner_id': project.owner_id,
                    'status': project.status,
                    'created_at': project.created_at.isoformat(),
                    'updated_at': project.updated_at.isoformat(),
                    'statistics': {
                        'documents': doc_count,
                        'analyses': analysis_count
                    }
                }
            }), 200
            
        finally:
            db.close()
            
    except Exception as e:
        return jsonify({'message': f'Failed to get project: {str(e)}'}), 500

@app.route('/api/projects/<project_id>', methods=['PUT'])
@token_required
@project_access_required
def update_project_api(project_id):
    """Update project details"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'message': 'No data provided'}), 400
        
        db = next(get_db())
        try:
            project = db.query(Project).filter(Project.id == project_id).first()
            
            if not project:
                return jsonify({'message': 'Project not found'}), 404
            
            # Update fields
            if 'name' in data:
                project.name = data['name']
            
            if 'description' in data:
                project.description = data['description']
            
            if 'status' in data:
                project.status = data['status']
            
            project.updated_at = datetime.now(timezone.utc)
            db.commit()
            
            return jsonify({
                'message': 'Project updated successfully',
                'project': {
                    'id': project.id,
                    'name': project.name,
                    'description': project.description,
                    'status': project.status,
                    'updated_at': project.updated_at.isoformat()
                }
            }), 200
            
        finally:
            db.close()
            
    except Exception as e:
        return jsonify({'message': f'Failed to update project: {str(e)}'}), 500

@app.route('/api/projects/<project_id>/members', methods=['POST'])
@token_required
@project_access_required
def add_project_member_api(project_id):
    """Add member to project"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'message': 'No data provided'}), 400
        
        user_email = data.get('user_email')
        role = data.get('role', ProjectUserRole.MEMBER)
        
        if not user_email:
            return jsonify({'message': 'User email is required'}), 400
        
        db = next(get_db())
        try:
            # Find user by email
            user = db.query(User).filter(User.email == user_email).first()
            
            if not user:
                return jsonify({'message': 'User not found'}), 404
            
            # Add user to project
            success = add_user_to_project(db, user.id, project_id, role)
            
            if success:
                return jsonify({
                    'message': 'User added to project successfully',
                    'user': {
                        'id': user.id,
                        'email': user.email,
                        'name': user.name,
                        'role': role
                    }
                }), 200
            else:
                return jsonify({'message': 'User is already a member of this project'}), 409
            
        finally:
            db.close()
            
    except Exception as e:
        return jsonify({'message': f'Failed to add member: {str(e)}'}), 500

# ============================================================================
# TEXT INPUT ENDPOINTS
# ============================================================================

@app.route('/api/projects/<project_id>/text-inputs', methods=['POST'])
@token_required
@project_access_required
def create_text_input_api(project_id):
    """Create text input for project"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'message': 'No data provided'}), 400
        
        title = data.get('title')
        content = data.get('content')
        template_id = data.get('template_id')
        
        if not title or not content:
            return jsonify({'message': 'Title and content are required'}), 400
        
        db = next(get_db())
        try:
            # Save text input
            text_input = save_text_input(
                db=db,
                title=title,
                content=content,
                user_id=request.current_user['id'],
                project_id=project_id,
                template_id=template_id
            )
            
            return jsonify({
                'message': 'Text input created successfully',
                'text_input': {
                    'id': text_input.id,
                    'title': text_input.title,
                    'content': text_input.content,
                    'created_at': text_input.created_at.isoformat(),
                    'version': text_input.version
                }
            }), 201
            
        finally:
            db.close()
            
    except Exception as e:
        return jsonify({'message': f'Failed to create text input: {str(e)}'}), 500

@app.route('/api/projects/<project_id>/text-inputs', methods=['GET'])
@token_required
@project_access_required
def get_text_inputs_api(project_id):
    """Get all text inputs for project"""
    try:
        db = next(get_db())
        try:
            text_inputs = db.query(TextInput).filter(
                TextInput.project_id == project_id,
                TextInput.is_latest == True
            ).all()
            
            return jsonify({
                'text_inputs': [
                    {
                        'id': ti.id,
                        'title': ti.title,
                        'content': ti.content,
                        'created_at': ti.created_at.isoformat(),
                        'updated_at': ti.updated_at.isoformat(),
                        'version': ti.version,
                        'user_id': ti.user_id
                    }
                    for ti in text_inputs
                ]
            }), 200
            
        finally:
            db.close()
            
    except Exception as e:
        return jsonify({'message': f'Failed to get text inputs: {str(e)}'}), 500

# ============================================================================
# ENHANCED INTEGRATION ENDPOINTS
# ============================================================================

# ============================================================================
# Helper Functions
# ============================================================================

def get_current_user_id():
    """Get current user ID from request context"""
    if hasattr(request, 'current_user') and request.current_user:
        return request.current_user.get('id')
    return None

# ============================================================================
# OneDrive Integration Endpoints
# ============================================================================

@app.route('/api/integrations/onedrive/status', methods=['GET'])
@token_required
def get_onedrive_status_api():
    """Get OneDrive connection status for current user"""
    try:
        # Get current user ID from token
        current_user_id = get_current_user_id()
        
        # Check if OneDrive is configured
        client_id = os.getenv('ONEDRIVE_CLIENT_ID')
        if not client_id:
            return jsonify({
                'connected': False,
                'message': 'OneDrive integration not configured',
                'configured': False
            }), 200
        
        # Check if user has connected their OneDrive
        user_token = integration_manager.user_onedrive_tokens.get(current_user_id)
        if user_token:
            # Check if token is still valid
            if time.time() < user_token.get('expires_at', 0):
                return jsonify({
                    'connected': True,
                    'message': 'OneDrive connected successfully',
                    'configured': True,
                    'user_connected': True
                }), 200
            else:
                # Token expired, remove it
                del integration_manager.user_onedrive_tokens[current_user_id]
        
        return jsonify({
            'connected': False,
            'message': 'OneDrive not connected for this user',
            'configured': True,
            'user_connected': False
        }), 200
        
    except Exception as e:
        return jsonify({'message': f'Failed to check OneDrive status: {str(e)}'}), 500

@app.route('/api/integrations/onedrive/auth', methods=['GET'])
@token_required
def get_onedrive_auth_url_api():
    """Get OneDrive authorization URL for current user"""
    try:
        current_user_id = get_current_user_id()
        
        # Generate state parameter for security
        state = f"user_{current_user_id}_{int(time.time())}"
        
        auth_url = integration_manager.get_onedrive_auth_url(current_user_id, state)
        
        if auth_url:
            return jsonify({
                'auth_url': auth_url,
                'state': state
            }), 200
        else:
            return jsonify({'message': 'Failed to generate auth URL'}), 500
            
    except Exception as e:
        return jsonify({'message': f'Failed to get auth URL: {str(e)}'}), 500

@app.route('/api/integrations/onedrive/callback', methods=['GET'])
def onedrive_callback_api():
    """Handle OneDrive authorization callback"""
    try:
        # Get parameters from callback
        auth_code = request.args.get('code')
        state = request.args.get('state')
        error = request.args.get('error')
        
        if error:
            return jsonify({'message': f'Authorization failed: {error}'}), 400
        
        if not auth_code or not state:
            return jsonify({'message': 'Missing authorization code or state'}), 400
        
        # Extract user ID from state
        try:
            user_id = state.split('_')[1]
        except:
            return jsonify({'message': 'Invalid state parameter'}), 400
        
        # Handle the callback
        success = integration_manager.handle_onedrive_callback(user_id, auth_code)
        
        if success:
            return jsonify({
                'message': 'OneDrive connected successfully! You can now close this window.',
                'success': True
            }), 200
        else:
            return jsonify({'message': 'Failed to connect OneDrive'}), 500
            
    except Exception as e:
        return jsonify({'message': f'Callback handling failed: {str(e)}'}), 500

@app.route('/api/integrations/onedrive/files', methods=['GET'])
@token_required
def get_onedrive_files_api():
    """Get OneDrive files for current user"""
    try:
        current_user_id = get_current_user_id()
        folder_id = request.args.get('folder_id')
        files = integration_manager.get_onedrive_files(folder_id, current_user_id)
        
        return jsonify({
            'files': files
        }), 200
        
    except Exception as e:
        return jsonify({'message': f'Failed to get OneDrive files: {str(e)}'}), 500

@app.route('/api/integrations/onedrive/download/<file_id>', methods=['GET'])
@token_required
def download_onedrive_file_api(file_id):
    """Download file from OneDrive for current user"""
    try:
        current_user_id = get_current_user_id()
        file_content = integration_manager.download_onedrive_file(file_id, current_user_id)
        
        if file_content is None:
            return jsonify({'message': 'Failed to download file from OneDrive'}), 500
        
        # Get file metadata to determine filename and type
        files = integration_manager.get_onedrive_files(user_id=current_user_id)
        file_info = next((f for f in files if f.get('id') == file_id), None)
        
        if not file_info:
            return jsonify({'message': 'File not found'}), 404
        
        filename = file_info.get('name', 'downloaded_file')
        file_type = file_info.get('file_type', 'application/octet-stream')
        
        # Create response with file content
        response = make_response(file_content)
        response.headers['Content-Type'] = file_type
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        return jsonify({'message': f'Failed to download file from OneDrive: {str(e)}'}), 500

@app.route('/api/integrations/onedrive/upload', methods=['POST'])
@token_required
def upload_to_onedrive_api():
    """Upload file to OneDrive"""
    try:
        if 'file' not in request.files:
            return jsonify({'message': 'No file provided'}), 400
        
        file = request.files['file']
        folder_id = request.form.get('folder_id')
        
        if file.filename == '':
            return jsonify({'message': 'No file selected'}), 400
        
        # Save file temporarily
        filename = secure_filename(file.filename)
        temp_path = os.path.join('uploads', filename)
        file.save(temp_path)
        
        try:
            # Upload to OneDrive
            result = integration_manager.upload_to_onedrive(temp_path, filename, folder_id)
            
            if result:
                return jsonify({
                    'message': 'File uploaded to OneDrive successfully',
                    'file': result
                }), 200
            else:
                return jsonify({'message': 'Failed to upload to OneDrive'}), 500
                
        finally:
            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
    except Exception as e:
        return jsonify({'message': f'Failed to upload to OneDrive: {str(e)}'}), 500

@app.route('/api/integrations/ado/connect', methods=['POST'])
@token_required
def connect_azure_devops_api():
    """Connect to Azure DevOps"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'message': 'No data provided'}), 400
        
        org_url = data.get('organization_url')
        pat_token = data.get('pat_token')
        
        if not org_url or not pat_token:
            return jsonify({'message': 'Organization URL and PAT token are required'}), 400
        
        # Setup connection
        success = integration_manager.setup_azure_devops(org_url, pat_token)
        
        if success:
            return jsonify({
                'message': 'Azure DevOps connected successfully'
            }), 200
        else:
            return jsonify({'message': 'Failed to connect to Azure DevOps'}), 500
        
    except Exception as e:
        return jsonify({'message': f'Failed to connect to Azure DevOps: {str(e)}'}), 500

@app.route('/api/integrations/ado/projects', methods=['GET'])
@token_required
def get_ado_projects_api():
    """Get Azure DevOps projects"""
    try:
        org_url = request.args.get('organization_url')
        
        if not org_url:
            return jsonify({'message': 'Organization URL is required'}), 400
        
        projects = integration_manager.get_ado_projects(org_url)
        
        return jsonify({
            'projects': projects
        }), 200
        
    except Exception as e:
        return jsonify({'message': f'Failed to get Azure DevOps projects: {str(e)}'}), 500

# ============================================================================
# ENHANCED HEALTH CHECK
# ============================================================================

@app.route('/api/health', methods=['GET'])
def health_check():
    """Enhanced health check endpoint"""
    try:
        # Check database connection
        db = next(get_db())
        db.execute(text("SELECT 1"))
        db.close()
        
        # Check integration manager status
        integration_status = integration_manager.get_status()
        
        return jsonify({
            'status': 'healthy',
            'timestamp': datetime.now(timezone.utc).isoformat(),
            'integrations': integration_status,
            'features': {
                'multi_user': True,
                'project_management': True,
                'document_analysis': True,
                'onedrive_integration': True,
                'azure_devops_integration': True
            }
        }), 200
        
    except Exception as e:
        return jsonify({
            'status': 'unhealthy',
            'error': str(e),
            'timestamp': datetime.now(timezone.utc).isoformat()
        }), 500

if __name__ == '__main__':
    print("🚀 Starting Enhanced BA Agent API...")
    print("✅ Multi-user authentication enabled")
    print("✅ Project management enabled")
    print("✅ Integration services enabled")
    print("✅ Document analysis enabled")
    print("✅ All original functionality preserved")
    
    app.run(debug=True, port=5000)
